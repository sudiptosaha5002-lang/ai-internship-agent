"""
AI Internship Finder Agent — Flask Backend
==========================================
A Flask server that powers the AI Internship Finder chatbot.
Uses LangChain + LangGraph with Google Gemini to create an intelligent agent that can
search for internships based on user skills, preferences, and resume.
"""

import os
import sys
# Force UTF-8 output so emoji in print() don't crash on Windows terminals
if hasattr(sys.stdout, 'reconfigure'):
    getattr(sys.stdout, 'reconfigure')(encoding='utf-8', errors='replace')
# BUILD_TAG: 2026-05-09-final-v1
import json
import tempfile
import base64
import io
import threading
from collections import OrderedDict
import requests
from flask import Flask, request, jsonify, render_template
from flask_cors import CORS
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from fpdf import FPDF
import docx as docx_lib
from email_validator import validate_email, EmailNotValidError

def _validate_email_deep(email):
    """
    Perform deep validation of an email:
    1. Syntax check
    2. DNS MX record check (does the domain exist and have a mail server?)
    """
    try:
        # check_deliverability=True performs the DNS MX record check
        valid = validate_email(email, check_deliverability=True, timeout=10)
        return True, valid.email
    except EmailNotValidError as e:
        return False, str(e)

from pdf2image import convert_from_bytes
import base64

from langchain_core.messages import HumanMessage, AIMessage, SystemMessage

def _safety_check_resume(resume_bytes):
    """Analyze resume for inappropriate images or content using Gemini Vision. Returns (is_safe, reason)."""
    llm_instance = get_llm()
    if not llm_instance:
        return True, "Safe (AI Skip)"
        
    try:
        # Only check the first page for speed and token efficiency
        images = convert_from_bytes(resume_bytes, first_page=1, last_page=1)
        if not images:
            return True, "No images"

        # Encode first page to base64
        import io
        img_byte_arr = io.BytesIO()
        images[0].save(img_byte_arr, format='JPEG', quality=80)
        base64_image = base64.b64encode(img_byte_arr.getvalue()).decode('utf-8')

        prompt = "Analyze this document image. Does it contain any inappropriate content, violent imagery, or random non-professional 'violated objects' that don't belong in a resume? Answer ONLY 'SAFE' or 'VIOLATED: <reason>'."
        
        # Multimodal message for Gemini
        message = HumanMessage(
            content=[
                {"type": "text", "text": prompt},
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
            ]
        )
        response = llm_instance.invoke([message])
        content = _normalize_llm_content(getattr(response, "content", "SAFE")).upper()
        
        if "VIOLATED" in content:
            return False, content
        return True, "Safe"
    except Exception as e:
        print(f"[SafetyCheck] Error: {e}")
        return True, "Safe (Error Fallback)"
from langgraph.prebuilt import create_react_agent

from tools import search_internships, filter_india_jobs
from mock_interview import mock_interview_bp
from email_utils import send_analysis_email

# ──────────────────────────────────────────────
# Load environment variables
# ──────────────────────────────────────────────
load_dotenv()

app = Flask(
    __name__,
    template_folder="internship_ai_agent",
    static_folder="internship_ai_agent",
    static_url_path="/static"
)
CORS(app)  # type: ignore
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024  # 32 MB max upload

@app.after_request
def add_header(response):
    """Add cache headers for static assets to improve load speed, but disable for HTML."""
    if request.path.endswith(('.html', '/')) or '/mock-interview' in request.path:
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
    elif request.path.startswith('/images/') or request.path.endswith(('.mp4', '.jpg', '.png', '.css', '.js')):
        response.cache_control.max_age = 604800 # 1 week
    return response

import uuid
def _resolve_upload_folder():
    explicit = os.getenv("UPLOAD_DIR", "").strip()
    if explicit:
        return explicit
    if os.getenv("RENDER") and os.path.isdir("/var/data"):
        return "/var/data/uploads"
    return os.path.join(os.path.dirname(__file__), 'uploads')


UPLOAD_FOLDER = _resolve_upload_folder()
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
RESUME_VARIATION_MEMORY_PATH = os.path.join(UPLOAD_FOLDER, 'resume_variation_memory.json')
RESUME_VARIATION_LOCK = threading.RLock()
CHAT_HISTORY_LOCK = threading.RLock()
MAX_CHAT_SESSIONS = max(50, int(os.getenv("MAX_CHAT_SESSIONS", "2000")))
MAX_HISTORY_MESSAGES = max(5, int(os.getenv("MAX_HISTORY_MESSAGES", "6")))
EXTERNAL_HTTP_TIMEOUT = max(3, int(os.getenv("EXTERNAL_HTTP_TIMEOUT", "20")))

# ── SQLite Database ──
import database as db

def get_resume_stream_from_req(req):
    # 1. Direct file upload takes priority
    if 'resume' in req.files and req.files['resume'].filename != '':
        file_obj = req.files['resume']
        payload = file_obj.read()
        file_obj.seek(0)
        return io.BytesIO(payload)

    # 2. Try identifying the session_id
    session_id = req.form.get("session_id") if req.form else None
    if not session_id and req.is_json:
        session_id = req.json.get("session_id")

    # 3. Check if the session_id from request has a file
    if session_id:
        filepath = os.path.join(UPLOAD_FOLDER, f"{session_id}.pdf")
        print(f"[DEBUG] Checking resume for SID {session_id} at {filepath}")
        if os.path.exists(filepath):
            print(f"[DEBUG] Found resume for SID {session_id}")
            with open(filepath, 'rb') as f:
                return io.BytesIO(f.read())
        else:
            print(f"[DEBUG] No file at {filepath}")

    # 4. If no file yet, try identifying via Auth Token (the reliable way for logged-in users)
    auth_header = req.headers.get('Authorization', '')
    if auth_header.startswith('Bearer '):
        token = auth_header.replace('Bearer ', '').strip()
        if token and token != 'null' and token != 'undefined':
            user = db.get_user_by_token(token)
            if user and user.get('session_id'):
                sid = user.get('session_id')
                filepath = os.path.join(UPLOAD_FOLDER, f"{sid}.pdf")
                print(f"[DEBUG] Checking resume via Token for SID {sid} at {filepath}")
                if os.path.exists(filepath):
                    print(f"[DEBUG] Found resume via Token for SID {sid}")
                    with open(filepath, 'rb') as f:
                        return io.BytesIO(f.read())
                else:
                    print(f"[DEBUG] No file at {filepath} (via token)")

    return None


def _load_resume_variation_memory():
    with RESUME_VARIATION_LOCK:
        try:
            if not os.path.exists(RESUME_VARIATION_MEMORY_PATH):
                return {}
            with open(RESUME_VARIATION_MEMORY_PATH, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return data if isinstance(data, dict) else {}
        except Exception:
            return {}


def _save_resume_variation_memory(data):
    with RESUME_VARIATION_LOCK:
        try:
            temp_path = f"{RESUME_VARIATION_MEMORY_PATH}.{uuid.uuid4().hex}.tmp"
            with open(temp_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            for attempt in range(5):
                try:
                    os.replace(temp_path, RESUME_VARIATION_MEMORY_PATH)
                    break
                except PermissionError:
                    if attempt == 4:
                        raise
                    import time
                    time.sleep(0.03 * (attempt + 1))
        except Exception as ex:
            print(f"[WARN] Failed to save resume variation memory: {ex}")


def _get_resume_variation_history(user_id):
    if not user_id:
        return {'summary': [], 'experience': []}
    memory = _load_resume_variation_memory()
    entry = memory.get(str(user_id), {})
    if not isinstance(entry, dict):
        return {'summary': [], 'experience': []}
    summary = entry.get('summary', [])
    experience = entry.get('experience', [])
    if not isinstance(summary, list):
        summary = []
    if not isinstance(experience, list):
        experience = []
    return {'summary': summary, 'experience': experience}


def _remember_resume_variation(user_id, summary_text, experience_text, keep=6):
    if not user_id:
        return
    with RESUME_VARIATION_LOCK:
        try:
            if not os.path.exists(RESUME_VARIATION_MEMORY_PATH):
                memory = {}
            else:
                with open(RESUME_VARIATION_MEMORY_PATH, 'r', encoding='utf-8') as f:
                    loaded = json.load(f)
                memory = loaded if isinstance(loaded, dict) else {}
        except Exception:
            memory = {}

        entry = memory.get(str(user_id), {})
        if not isinstance(entry, dict):
            entry = {}
        summary_hist = entry.get('summary', [])
        experience_hist = entry.get('experience', [])
        if not isinstance(summary_hist, list):
            summary_hist = []
        if not isinstance(experience_hist, list):
            experience_hist = []

        if summary_text and summary_text.strip():
            summary_hist.append(summary_text.strip())
        if experience_text and experience_text.strip():
            experience_hist.append(experience_text.strip())

        entry['summary'] = summary_hist[-keep:]
        entry['experience'] = experience_hist[-keep:]
        memory[str(user_id)] = entry

        temp_path = f"{RESUME_VARIATION_MEMORY_PATH}.{uuid.uuid4().hex}.tmp"
        try:
            with open(temp_path, 'w', encoding='utf-8') as f:
                json.dump(memory, f, ensure_ascii=False, indent=2)
            for attempt in range(5):
                try:
                    os.replace(temp_path, RESUME_VARIATION_MEMORY_PATH)
                    break
                except PermissionError:
                    if attempt == 4:
                        raise
                    import time
                    time.sleep(0.03 * (attempt + 1))
        except Exception as ex:
            print(f"[WARN] Failed to save resume variation memory: {ex}")

app.register_blueprint(mock_interview_bp)

# ──────────────────────────────────────────────
# Native Auth API Routes (no Node.js needed)
# ──────────────────────────────────────────────
import hashlib
import secrets

def _hash_password(pw):
    return hashlib.sha256(pw.encode()).hexdigest()

def _make_token(user_id):
    return secrets.token_hex(32)


@app.route('/api/auth/register', methods=['POST'])
def auth_register():
    """Handle registration from the main register.html form (first_name + last_name fields)."""
    try:
        # register.html sends form-data with first_name / last_name
        first_name = request.form.get('first_name', '').strip()
        last_name  = request.form.get('last_name', '').strip()
        email      = request.form.get('email', '').strip().lower()
        password   = request.form.get('password', '')
        target_role = request.form.get('role', '').strip()
        full_name  = (first_name + ' ' + last_name).strip() or email.split('@')[0]

        if not email or not password:
            return jsonify({'status': 'error', 'error': 'Email and password are required'}), 400

        # Deep validation: Check if email domain exists
        is_valid, validation_msg = _validate_email_deep(email)
        if not is_valid:
            return jsonify({'status': 'error', 'error': f"Invalid email: Please provide an existing email address. ({validation_msg})"}), 400

        # Check if user already exists
        existing_user = db.get_user_by_email(email)
        if existing_user:
            return jsonify({'status': 'error', 'error': 'ALready exists this email login through the sign in page'}), 409

        # Handle optional resume upload
        skills = []
        role   = target_role or 'Software Engineer'
        session_id = None
        if 'resume' in request.files:
            resume_file = request.files['resume']
            resume_bytes = resume_file.read()
            resume_file.seek(0)

            # --- SAFETY CHECK ---
            is_safe, safety_reason = _safety_check_resume(resume_bytes)
            if not is_safe:
                return jsonify({'status': 'error', 'error': f"SECURITY ALERT: Your resume contains inappropriate visual content ({safety_reason}). Registration blocked."}), 403

            session_id  = str(uuid.uuid4())
            filepath    = os.path.join(UPLOAD_FOLDER, f"{session_id}.pdf")
            resume_file.save(filepath)
            try:
                from mock_interview import extract_text_from_pdf, extract_skills, infer_role
                with open(filepath, 'rb') as f:
                    text = extract_text_from_pdf(f)
                skills = extract_skills(text)
                if skills: role = infer_role(skills)
            except: pass

        user_data = db.create_user(
            email=email,
            password_hash=_hash_password(password),
            full_name=full_name,
            target_role=role,
            session_id=session_id or '',
            extracted_skills=skills,
            extracted_role=role,
            is_verified=1
        )

        # Auto-login after registration
        if not user_data:
            raise ValueError("Failed to create user")
        token = db.refresh_user_token(user_data['_id'])
        return jsonify({
            'token': token, 
            'session_id': user_data.get('session_id', ''),
            'user': db.user_to_safe_dict(user_data),
            'skills': skills,
            'role': role
        }), 201
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'status': 'error', 'error': str(e)}), 500


@app.route('/api/auth/signup', methods=['POST'])
def auth_signup():
    try:
        # Accept both JSON and form-data
        if request.is_json:
            data = request.get_json()
            email = data.get('email', '').strip().lower()
            password = data.get('password', '')
            full_name = data.get('fullName', data.get('full_name', ''))
        else:
            email = request.form.get('email', '').strip().lower()
            password = request.form.get('password', '')
            full_name = request.form.get('fullName', request.form.get('full_name', ''))

        if not email or not password:
            return jsonify({"error": "Email and password are required"}), 400

        # Check if user already exists
        existing_user = db.get_user_by_email(email)
        if existing_user:
            return jsonify({"error": "ALready exists this email login through the sign in page"}), 409

        # Handle optional resume upload
        skills = []
        role = 'Software Engineer'
        session_id = None
        if 'resume' in request.files:
            resume_file = request.files['resume']
            resume_bytes = resume_file.read()
            resume_file.seek(0)

            # --- SAFETY CHECK ---
            is_safe, safety_reason = _safety_check_resume(resume_bytes)
            if not is_safe:
                return jsonify({'error': f"SECURITY ALERT: Inappropriate visual content detected ({safety_reason}). Account creation blocked."}), 403

            session_id = str(uuid.uuid4())
            filepath = os.path.join(UPLOAD_FOLDER, f"{session_id}.pdf")
            resume_file.save(filepath)
            try:
                from mock_interview import extract_text_from_pdf, extract_skills, infer_role
                with open(filepath, 'rb') as f:
                    text = extract_text_from_pdf(f)
                skills = extract_skills(text)
                role = infer_role(skills) if skills else 'Software Engineer'
            except: pass

        user_data = db.create_user(
            email=email,
            password_hash=_hash_password(password),
            full_name=full_name or email.split('@')[0],
            target_role=role,
            session_id=session_id or '',
            extracted_skills=skills,
            extracted_role=role,
            is_verified=1
        )

        if not user_data:
            raise ValueError("Failed to create user")
        token = db.refresh_user_token(user_data['_id'])
        return jsonify({
            'token': token,
            'session_id': user_data.get('session_id', ''),
            'user': db.user_to_safe_dict(user_data),
            'skills': skills,
            'role': role
        }), 201
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/auth/login', methods=['POST'])
def auth_login():
    try:
        if request.is_json:
            data = request.get_json()
            email = data.get('email', '').strip().lower()
            password = data.get('password', '')
        else:
            email = request.form.get('email', '').strip().lower()
            password = request.form.get('password', '')

        if not email or not password:
            return jsonify({'error': 'Email and password are required'}), 400

        # Deep validation: Check if email domain exists
        is_valid, validation_msg = _validate_email_deep(email)
        if not is_valid:
            return jsonify({'error': f"Please provide a valid, existing email address. ({validation_msg})"}), 400

        user = db.get_user_by_email(email)
        if not user:
            # Auto-create verified user
            user_data = db.create_user(
                email=email,
                password_hash=_hash_password(password),
                full_name=email.split('@')[0],
                target_role='Software Engineer',
                is_verified=1
            )
            if not user_data:
                raise ValueError("Failed to create user")
            token = db.refresh_user_token(user_data['_id'])
            return jsonify({
                'token': token,
                'session_id': user_data.get('session_id', ''),
                'user': db.user_to_safe_dict(user_data),
                'skills': [],
                'role': 'Software Engineer'
            }), 200

        # Check if banned
        if user.get('is_banned'):
            return jsonify({'error': 'ACCESS DENIED: Your account is permanently banned for violating our community safety standards.'}), 403

        # Verify password
        if user.get('password_hash') != _hash_password(password):
            return jsonify({'error': 'Invalid email or password'}), 401

        # Check verification status (fallback for older accounts)
        if not user.get('is_verified'):
            db.mark_user_verified(email)
            user['is_verified'] = 1

        # Refresh token on every login
        token = db.refresh_user_token(user['_id'])

        skills    = user.get('extractedSkills', [])
        role      = user.get('extractedRole', user.get('extracted_role', 'Software Engineer'))
        
        safe_user = db.user_to_safe_dict(user)
        safe_user['token'] = token
        return jsonify({'token': token, 'session_id': user.get('session_id', ''), 'user': safe_user, 'skills': skills, 'role': role})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/auth/verify-otp', methods=['POST'])
def auth_verify_otp():
    try:
        data = request.get_json() or {}
        email = data.get('email')
        otp = data.get('otp')
        
        if not email or not otp:
            return jsonify({'error': 'Email and OTP are required'}), 400
            
        if db.verify_otp(email, otp):
            db.mark_user_verified(email)
            user = db.get_user_by_email(email)
            if not user:
                return jsonify({'error': 'User not found'}), 404
            return jsonify({
                'status': 'success',
                'message': 'Email verified successfully!',
                'token': user['token'],
                'session_id': user.get('session_id', ''),
                'user': db.user_to_safe_dict(user)
            })
        else:
            return jsonify({'error': 'Invalid or expired OTP code'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/auth/resend-otp', methods=['POST'])
def auth_resend_otp():
    try:
        data = request.get_json() or {}
        email = data.get('email')
        if not email:
            return jsonify({'error': 'Email is required'}), 400
            
        import random
        otp = f"{random.randint(100000, 999999)}"
        db.save_otp(email, otp)
        from email_utils import send_otp_email
        send_otp_email(email, otp)
        
        return jsonify({'status': 'success', 'message': 'New OTP sent to your inbox.'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/auth/forgot-password', methods=['POST'])
def auth_forgot_password():
    try:
        data = request.get_json() or {}
        email = data.get('email', '').strip().lower()
        if not email:
            return jsonify({'error': 'Email is required'}), 400
            
        user = db.get_user_by_email(email)
        if not user:
            # We return success even if user doesn't exist for security (avoid enumeration)
            # but in a friendly app we can just say "If account exists, email sent"
            return jsonify({'status': 'success', 'message': 'If an account exists with this email, a reset code has been sent.'})

        import random
        otp = f"{random.randint(100000, 999999)}"
        db.save_reset_otp(email, otp)
        from email_utils import send_reset_email
        send_reset_email(email, otp)
        
        return jsonify({'status': 'success', 'message': 'Reset code sent to your inbox.'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/auth/reset-password', methods=['POST'])
def auth_reset_password():
    try:
        data = request.get_json() or {}
        email = data.get('email', '').strip().lower()
        otp = data.get('otp', '').strip()
        new_password = data.get('password', '')

        if not email or not otp or not new_password:
            return jsonify({'error': 'Email, OTP, and new password are required'}), 400
            
        if db.verify_reset_otp(email, otp):
            db.update_user_password(email, _hash_password(new_password))
            user = db.get_user_by_email(email)
            if not user:
                return jsonify({'error': 'User not found'}), 404
            token = db.refresh_user_token(user['_id'])
            return jsonify({
                'status': 'success', 
                'message': 'Password reset successful!',
                'token': token,
                'session_id': user.get('session_id', ''),
                'user': db.user_to_safe_dict(user)
            })
        else:
            return jsonify({'error': 'Invalid or expired reset code'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/auth/upload-resume', methods=['POST'])
def auth_upload_resume():
    try:
        auth_header = request.headers.get('Authorization', '')
        token = auth_header.replace('Bearer ', '').strip()

        user = db.get_user_by_token(token)
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401

        if 'resume' not in request.files:
            return jsonify({'error': 'No resume file provided'}), 400

        resume_file = request.files['resume']
        resume_bytes = resume_file.read()
        resume_file.seek(0)

        # --- SAFETY CHECK ---
        is_safe, safety_reason = _safety_check_resume(resume_bytes)
        if not is_safe:
            # BAN THE USER IMMEDIATELY
            db.update_user(user['_id'], is_banned=1)
            return jsonify({'error': f"SECURITY ALERT: Inappropriate content detected ({safety_reason}). Your account has been permanently BANNED."}), 403

        session_id = str(uuid.uuid4())
        filepath = os.path.join(UPLOAD_FOLDER, f"{session_id}.pdf")
        resume_file.save(filepath)

        skills = []
        role = 'Software Engineer'
        full_name = None
        try:
            from mock_interview import extract_text_from_pdf, extract_skills, infer_role, extract_resume_profile
            with open(filepath, 'rb') as f:
                text = extract_text_from_pdf(f)
            skills = extract_skills(text)
            role = infer_role(skills) if skills else 'Software Engineer'
            
            # Extract profile info like name
            profile_data = extract_resume_profile(text)
            if profile_data and profile_data.get('fullName') and profile_data.get('fullName') != "[Candidate Full Name]":
                full_name = profile_data.get('fullName')
        except Exception as e:
            print(f"[UploadResume] Error: {e}")

        # Update user record
        update_args = {
            'extracted_skills': skills,
            'extracted_role': role,
            'session_id': session_id
        }
        if full_name:
            update_args['full_name'] = full_name
            
        db.update_user(user['_id'], **update_args)

        # Re-fetch updated user
        user = db.get_user_by_id(user['_id'])
        safe_user = db.user_to_safe_dict(user)
        return jsonify({'user': safe_user, 'skills': skills, 'role': role})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ──────────────────────────────────────────────
# Forgot Password Pages
# ──────────────────────────────────────────────

@app.route('/forgot-password')
def forgot_password_page():
    return render_template('forgot-password.html')


# ──────────────────────────────────────────────
# Profile & Interview History API Routes
# ──────────────────────────────────────────────

def _get_user_from_request():
    """Extract the authenticated user from the request (via Authorization header or session_id)."""
    auth_header = request.headers.get('Authorization', '')
    token = auth_header.replace('Bearer ', '').strip()
    if token:
        return db.get_user_by_token(token)
    return None


@app.route('/api/profile', methods=['GET'])
def api_get_profile():
    """Fetch the authenticated user's profile from the database."""
    try:
        user = _get_user_from_request()
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401

        profile = db.get_profile(user['_id'])
        # Merge user-level fields into profile for convenience
        profile['fullName'] = user.get('fullName', user.get('full_name', ''))
        profile['email'] = user.get('email', '')
        profile['role'] = user.get('extractedRole', user.get('extracted_role', ''))
        profile['extractedSkills'] = user.get('extractedSkills', [])
        return jsonify({'profile': profile})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/profile', methods=['POST'])
def api_save_profile():
    """Save or update the authenticated user's profile."""
    try:
        user = _get_user_from_request()
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401

        if request.is_json:
            data = request.get_json()
        else:
            data = request.form.to_dict()

        # Update the user's name if provided
        name = data.get('name', data.get('fullName', ''))
        if name:
            db.update_user(user['_id'], full_name=name)

        # Save profile fields
        db.save_profile(user['_id'], data)

        return jsonify({'status': 'success', 'message': 'Profile saved'})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/interview-history', methods=['GET'])
def api_get_interview_history():
    """Fetch the authenticated user's interview history."""
    try:
        user = _get_user_from_request()
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401

        history = db.get_interview_history(user['_id'])
        return jsonify({'history': history})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/interview-history', methods=['POST'])
def api_save_interview_history():
    """Save an interview session result."""
    try:
        user = _get_user_from_request()
        user_id = user['_id'] if user else None

        data = request.get_json() or {}
        session_id = db.save_interview_session(user_id, data)

        return jsonify({'status': 'success', 'session_id': session_id})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ──────────────────────────────────────────────
# LangChain Agent Setup
# ──────────────────────────────────────────────

# ── LLM Setup: Google Gemini 2.5 Flash (free tier) primary, graceful fallback ──
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY", "")
_llm_instance = None
_llm_resume_instance = None

def get_llm():
    global _llm_instance
    if _llm_instance is not None:
        return _llm_instance
    if not GOOGLE_API_KEY:
        return None
    try:
        from langchain_google_genai import ChatGoogleGenerativeAI
        _llm_instance = ChatGoogleGenerativeAI(
            model="gemini-2.5-flash",
            google_api_key=GOOGLE_API_KEY,
            temperature=0.4,
            convert_system_message_to_human=True,
            top_p=0.95,
            top_k=40
        )
        print("[OK] LLM: Google Gemini 2.5 Flash")
        return _llm_instance
    except Exception as _llm_err:
        print(f"[WARN] Failed to initialise Gemini LLM: {_llm_err}")
        return None

def get_llm_resume():
    global _llm_resume_instance
    if _llm_resume_instance is not None:
        return _llm_resume_instance
    if not GOOGLE_API_KEY:
        return None
    try:
        from langchain_google_genai import ChatGoogleGenerativeAI
        _llm_resume_instance = ChatGoogleGenerativeAI(
            model="gemini-2.5-flash",
            google_api_key=GOOGLE_API_KEY,
            temperature=0.8,
            convert_system_message_to_human=True,
            top_p=0.98,
            top_k=50
        )
        return _llm_resume_instance
    except Exception as _llm_err:
        print(f"[WARN] Failed to initialise Gemini Resume LLM: {_llm_err}")
        return None
if not GOOGLE_API_KEY:
    print("[WARN] GOOGLE_API_KEY not set - AI chat features will be unavailable.")
    print("       Set GOOGLE_API_KEY in your Modal secret or .env file.")
    print("       Get a free Gemini key at: https://aistudio.google.com/")
tools = [search_internships]

# System prompt for the AI agent (Optimized for speed & anti-hallucination)
SYSTEM_PROMPT = """You are InternshipAgent. Your goal is to help students find internships efficiently.

RULES:
1. ALWAYS use the `search_internships` tool to find internships.
2. The `search_internships` tool will return a fully formatted ```internship_cards``` block. YOU MUST ECHO THIS EXACT BLOCK TO THE USER VERBATIM!
3. DO NOT change the JSON. DO NOT change the apply_links. DO NOT invent URLs. DO NOT hallucinate.
4. Keep all conversational text extremely concise. Answer in 1 short sentence.
"""

# ──────────────────────────────────────────────
# In-memory chat history (per session — resets on server restart)
# ──────────────────────────────────────────────
chat_histories = OrderedDict()


def get_chat_history(session_id: str) -> list:
    """Get or create chat history for a session."""
    sid = session_id or "default"
    with CHAT_HISTORY_LOCK:
        history = chat_histories.get(sid)
        if history is None:
            history = [SystemMessage(content=SYSTEM_PROMPT)]
            chat_histories[sid] = history
        else:
            chat_histories.move_to_end(sid)

        while len(chat_histories) > MAX_CHAT_SESSIONS:
            chat_histories.popitem(last=False)
        return history


def append_chat_messages(session_id: str, messages):
    """Append messages to chat history and trim to bounded length."""
    if not isinstance(messages, list):
        messages = [messages]
    with CHAT_HISTORY_LOCK:
        history = get_chat_history(session_id)
        history.extend(messages)
        if len(history) > MAX_HISTORY_MESSAGES:
            history[:] = [history[0]] + history[-(MAX_HISTORY_MESSAGES - 1):]


def get_chat_history_snapshot(session_id: str) -> list:
    """Return a shallow copy so model calls aren't affected by concurrent writes."""
    with CHAT_HISTORY_LOCK:
        return list(get_chat_history(session_id))


# ──────────────────────────────────────────────
# Routes
# ──────────────────────────────────────────────

@app.route("/")
def index():
    """Serve the premium liquid glass landing page."""
    return render_template("landing.html")

@app.route("/about-us")
def about_us():
    """Serve the about us page."""
    return render_template("about-us.html")

@app.route("/privacy")
@app.route("/privacy-policy")
def privacy():
    """Serve the privacy policy page."""
    return render_template("privacy-policy.html")

@app.route("/terms-and-conditions")
def terms_and_conditions():
    """Serve the terms and conditions page."""
    return render_template("terms-and-conditions.html")

@app.route("/contact")
@app.route("/contact-us")
def contact():
    """Serve the contact page."""
    return render_template("contact-us.html")

@app.route("/profile")
def profile():
    """Serve the user profile page."""
    return render_template("profile.html")

@app.route("/images/<path:filename>")
def serve_image(filename):
    """Serve images from the internship_ai_agent/images folder."""
    from flask import send_from_directory
    images_folder = os.path.join(os.path.dirname(__file__), 'internship_ai_agent', 'images')
    return send_from_directory(images_folder, filename)

@app.route("/user-dashboard")
def user_dashboard():
    """Serve the post-auth user dashboard."""
    return render_template("user-dashboard.html")

@app.route("/dashboard")
def dashboard():
    return render_template("user-dashboard.html")

@app.route("/mock-interview/", defaults={"path": ""})
@app.route("/mock-interview/<path:path>")
def serve_mock_interview(path):
    return render_template("mock-interview.html")

@app.route("/documentation")
def documentation():
    return render_template("documentation.html")

@app.route("/bot")
def bot():
    """Serve the original interactive chatbot UI."""
    return render_template("bot.html")

@app.route('/login')
def login():
    return render_template('login.html')

@app.route('/register')
def register():
    return render_template('register.html')

@app.route('/verify')
def verify():
    return render_template('verify.html')

# ── Auth routes are now proxied to Node.js backend via /api/auth/<path:subpath> ──

@app.route('/api/generate-resume', methods=['POST'])
def generate_resume():
    try:
        user = _get_user_from_request()
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401

        data = request.get_json() or {}
        variation_nonce = uuid.uuid4().hex[:10]
        variation_history = _get_resume_variation_history(user.get('_id'))

        def build_avoid_phrases(history_items, limit=8):
            import re
            phrases = []
            for item in history_items[-3:]:
                if not item or not isinstance(item, str):
                    continue
                for chunk in re.split(r'[.\n;]+', item):
                    cleaned = " ".join(chunk.strip().split())
                    if len(cleaned.split()) >= 4:
                        phrases.append(cleaned)
            unique = []
            seen = set()
            for phrase in phrases:
                key = phrase.lower()
                if key in seen:
                    continue
                seen.add(key)
                unique.append(phrase[:120])
                if len(unique) >= limit:
                    break
            return unique

        summary_avoid_phrases = build_avoid_phrases(variation_history.get('summary', []))
        experience_avoid_phrases = build_avoid_phrases(variation_history.get('experience', []))

        def format_avoid_phrases(phrases):
            if not phrases:
                return "None"
            return "\n".join(f'- "{p}"' for p in phrases)

        def pick_non_repeating(options, context_key):
            import random
            if not options:
                return ""
            history = variation_history.get(context_key, [])
            previous = history[-1].strip().lower() if history else ""
            filtered = [opt for opt in options if opt and opt.strip().lower() != previous]

            avoid_phrases = summary_avoid_phrases if context_key == 'summary' else experience_avoid_phrases
            phrase_filtered = [
                opt for opt in filtered
                if not any(p.lower() in opt.lower() for p in avoid_phrases[:3])
            ]
            pool = phrase_filtered or filtered or options
            return random.choice(pool)
        # FPDF Sanitization: Remove non-Latin-1 characters to prevent UnicodeEncodeError
        def sanitize_for_fpdf(text):
            if not text: return ""
            return text.encode('latin-1', 'ignore').decode('latin-1')

        # Professional FPDF Generation
        class ProResumePDF(FPDF):
            def header(self):
                pass
            def footer(self):
                self.set_y(-15)
                self.set_font('Arial', 'I', 8)
                self.set_text_color(150, 150, 150)
                self.cell(0, 10, 'Generated by InternAI Pro', 0, 0, 'C')

        pdf = ProResumePDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Colors
        accent_r, accent_g, accent_b = 33, 50, 94 # Dark Slate Blue
        text_r, text_g, text_b = 50, 50, 50
        
        # 1. HEADER SECTON (Colored background)
        pdf.set_fill_color(accent_r, accent_g, accent_b)
        pdf.rect(0, 0, 210, 45, 'F')
        
        # Name
        pdf.set_y(10)
        pdf.set_font('Arial', 'B', 26)
        pdf.set_text_color(255, 255, 255)
        name = f"{data.get('firstName', '')} {data.get('lastName', '')}".strip()
        if not name: name = user.get('full_name') or 'Professional Name'
        pdf.cell(0, 10, sanitize_for_fpdf(name.upper()), ln=True, align='C')
        
        # Role
        pdf.set_font('Arial', '', 14)
        pdf.set_text_color(200, 220, 255)
        pdf.cell(0, 8, sanitize_for_fpdf(data.get('profession', 'Software Engineer').upper()), ln=True, align='C')
        
        # Contact Info
        pdf.set_font('Arial', '', 10)
        pdf.set_text_color(255, 255, 255)
        contact = f"{data.get('location', '')}  |  {data.get('phone', '')}  |  {data.get('email', '')}"
        pdf.cell(0, 6, sanitize_for_fpdf(contact), ln=True, align='C')
        
        pdf.set_y(55) # Move below header
        
        # Helper function for sections
        def add_section_title(title):
            pdf.ln(4)
            pdf.set_font('Arial', 'B', 14)
            pdf.set_text_color(accent_r, accent_g, accent_b)
            pdf.cell(0, 8, sanitize_for_fpdf(title.upper()), ln=True)
            # Horizontal line
            y = pdf.get_y()
            pdf.set_draw_color(accent_r, accent_g, accent_b)
            pdf.line(10, y, 200, y)
            pdf.ln(3)

        # AI Enhancement of Text using Gemini 2.5 Flash with specialized prompts
        def enhance_text(text, context, user_info=None):
            """Aggressive AI enhancement with VARIED output each time."""
            from langchain_core.messages import HumanMessage
            import random
            
            user_info = user_info or {}
            profession = user_info.get('profession', 'Software Engineer')
            skills = user_info.get('skills', '')
            years_exp = user_info.get('years_exp', '2')
            
            # If text is empty or very minimal, generate from scratch
            if not text or len(text.strip()) < 5:
                if context == "professional summary":
                    return generate_professional_summary(user_info)
                elif context == "work experience description/achievements":
                    return generate_job_description(user_info)
                return text
            
            # VARIED PROMPTS - Different focus each time for unique output
            if context == "professional summary":
                variation_styles = [
                    {
                        "name": "Impact-Driven",
                        "start_words": ["Results-oriented", "Impact-driven", "Proven"],
                        "focus": "business outcomes and measurable results"
                    },
                    {
                        "name": "Innovation-Focused",
                        "start_words": ["Innovative", "Forward-thinking", "Strategic"],
                        "focus": "technical innovation and cutting-edge solutions"
                    },
                    {
                        "name": "Leadership-Oriented",
                        "start_words": ["Dynamic", "Visionary", "Driven"],
                        "focus": "team leadership and collaborative achievements"
                    }
                ]
                
                style = random.choice(variation_styles)
                start_word = random.choice(style["start_words"])
                
                prompt = f"""REWRITE this into a PROFESSIONAL resume summary with focus on {style["focus"]}.

INPUT: {text}
VARIATION_TOKEN: {variation_nonce}
RECENT_PHRASES_TO_AVOID:
{format_avoid_phrases(summary_avoid_phrases)}

STYLE: {style["name"]}
- Start with: "{start_word}"
- Emphasize: {style["focus"]}

REQUIREMENTS:
1. Exactly 3-4 sentences (different order/structure than previous versions)
2. MUST start with: "{start_word}"
3. Include specific years or domain expertise
4. Use strong verbs: Engineered, Led, Delivered, Architected, Optimized, Transformed
5. Show concrete business impact
6. NO personal pronouns (I, me, we)
7. NO generic phrases
8. Avoid reusing wording from RECENT_PHRASES_TO_AVOID

VARY THE STRUCTURE - Make it unique:
- Option A: [Adjective] [Profession] → Specialized in → Proven track record of → Dedicated to
- Option B: [Adjective] [Professional] → Track record of → Deep expertise in → Passionate about
- Option C: [Adjective] [Role] → Demonstrated success in → Technical excellence in → Committed to

Pick a DIFFERENT structure than usual. Output ONLY the summary - NO PREAMBLE."""

            elif context == "work experience description/achievements":
                variation_styles = [
                    {
                        "name": "Technical Excellence",
                        "focus": "technical depth and architectural decisions",
                        "lead_verbs": ["Engineered", "Architected", "Designed", "Implemented"]
                    },
                    {
                        "name": "Business Impact",
                        "focus": "revenue, cost savings, and business metrics",
                        "lead_verbs": ["Delivered", "Optimized", "Accelerated", "Transformed"]
                    },
                    {
                        "name": "Leadership & Scale",
                        "focus": "team leadership and scaling systems",
                        "lead_verbs": ["Led", "Spearheaded", "Pioneered", "Established"]
                    }
                ]
                
                style = random.choice(variation_styles)
                selected_verbs = ", ".join(random.sample(style["lead_verbs"], 2))
                
                prompt = f"""Transform this into PROFESSIONAL achievement bullets with focus on {style["focus"]}.

INPUT: {text}
VARIATION_TOKEN: {variation_nonce}
RECENT_PHRASES_TO_AVOID:
{format_avoid_phrases(experience_avoid_phrases)}
STYLE: {style["name"]}
PRIMARY VERBS: {selected_verbs}

REQUIREMENTS FOR EACH BULLET:
1. Start with: {selected_verbs} (or similar power verbs)
2. MUST have: Specific metrics/numbers with context (%, time, $, users, improvement)
3. Include: Technology/tool names
4. Format: [Verb] [What] using [Tech], [Result with number]
5. One line, scannable
6. NO gerunds, NO personal pronouns
7. Avoid reusing wording from RECENT_PHRASES_TO_AVOID

CREATE 4 VARIED BULLETS showing:
- Different metrics (% improvement, time saved, volume processed, cost reduction)
- Different technologies mentioned
- Different business/technical outcomes
- Different team/scope sizes

Examples of VARIED formats:
• "Engineered X using Y, improving Z by 45%" 
• "Led team of X, architecting Y that handles Z daily"
• "Optimized X, reducing Z from Y to Y (60% improvement)"
• "Delivered X using Y, enabling Z outcome"

Output ONLY bullets on separate lines starting with "- ", NO PREAMBLE."""


            else:
                prompt = f"Rewrite this professionally for a resume: {text}"
            
            try:
                print(f"[DEBUG] Calling LLM for {context} enhancement...")
                llm = get_llm_resume()
                if llm is None:
                    raise ValueError("LLM instance is not available")
                res = llm.invoke([HumanMessage(content=prompt)])
                content = res.content if isinstance(res.content, str) else str(res.content)
                result = content.strip()
                
                # Clean up formatting
                result = result.replace('```', '').replace('**', '').strip()
                
                # Validate output
                if result and len(result.strip()) > 10:
                    print(f"[SUCCESS] Enhanced {context}: {result[:80]}...")
                    return result
                else:
                    print(f"[WARN] LLM returned empty result for {context}, using fallback")
                    return fallback_enhance(text, context, user_info)
                    
            except Exception as e:
                print(f"[ERROR] LLM Enhancement failed ({context}): {e}")
                return fallback_enhance(text, context, user_info)
        
        def fallback_enhance(text, context, user_info):
            """Fallback enhancement with VARIED templates - never the same twice."""
            
            user_info = user_info or {}
            profession = user_info.get('profession', 'Software Engineer')
            skills = user_info.get('skills', '')
            years_exp = user_info.get('years_exp', '2')
            
            if context == "professional summary":
                kw_list = [s.strip() for s in skills.split(',')[:3]] if skills else ['modern technologies', 'scalable systems', 'innovative solutions']
                has_years = any(str(i) in text for i in range(1, 50))
                years = years_exp if has_years else "2+"
                
                # Multiple varied templates so fallback is never identical
                templates = [
                    f"Results-oriented {profession} with {years} years of proven track record in delivering impactful solutions. Deep expertise in {kw_list[0]} combined with strong technical foundation. Committed to driving innovation and leveraging technology to solve complex business challenges.",
                    
                    f"Driven {profession} with {years}+ years of hands-on experience architecting and delivering scalable systems. Specialized in {kw_list[0]} with demonstrated ability to optimize performance and drive measurable outcomes. Passionate about continuous learning and collaborating across teams to achieve organizational goals.",
                    
                    f"Innovative {profession} with {years} years of proven expertise in building solutions using {kw_list[0]}. Strong track record of translating complex requirements into technical excellence. Dedicated to leveraging best practices and emerging technologies to deliver competitive advantage and business value.",
                    
                    f"Strategic {profession} with {years} years of experience delivering high-impact solutions. Specialized in {kw_list[0]} with proven ability to optimize systems and enhance operational efficiency. Committed to fostering collaborative environments and driving technical excellence across projects."
                ]
                return pick_non_repeating(templates, 'summary')
            
            elif context == "work experience description/achievements":
                kw_list = [s.strip() for s in skills.split(',')[:3]] if skills else ['technical solutions', 'modern frameworks', 'scalable systems']
                
                # Multiple varied bullet templates
                template_sets = [
                    [
                        f"- Engineered robust {kw_list[0]} using modern technologies, improving system performance by 35%+",
                        "- Led cross-functional initiatives resulting in 40% efficiency gain across team productivity",
                        f"- Architected and implemented {kw_list[1]} reducing infrastructure costs by 25%",
                        "- Mentored junior developers and established code quality standards improving delivery velocity"
                    ],
                    [
                        f"- Delivered enterprise-grade {kw_list[0]} handling 10K+ concurrent users with 99.9% availability",
                        "- Optimized database operations and API performance, reducing latency by 60%",
                        f"- Spearheaded migration to {kw_list[1]}, enabling faster feature releases and deployments",
                        "- Collaborated with product teams to define technical roadmap and deliver on aggressive timelines"
                    ],
                    [
                        f"- Architected {kw_list[0]} using {kw_list[1]}, increasing throughput by 50% and reducing costs 30%",
                        "- Led team of 3+ engineers developing microservices handling millions of requests daily",
                        "- Implemented comprehensive testing and CI/CD pipelines reducing deployment time by 75%",
                        "- Drove technical innovation initiatives resulting in improved code maintainability and team efficiency"
                    ],
                    [
                        f"- Engineered high-performance {kw_list[0]} improving user experience and engagement by 45%",
                        "- Optimized cloud infrastructure reducing monthly operational expenses by $40K+",
                        f"- Led redesign of {kw_list[1]} architecture improving scalability for 5x user growth",
                        "- Established best practices and mentoring program improving team capabilities and retention"
                    ]
                ]
                chosen_set = pick_non_repeating(["\n".join(s) for s in template_sets], 'experience')
                return chosen_set
            
            return text
        
        def generate_professional_summary(user_info):
            """Generate compelling professional summary - VARIED each time."""
            from langchain_core.messages import HumanMessage
            import random
            
            profession = user_info.get('profession', 'Software Engineer')
            skills = user_info.get('skills', '')
            location = user_info.get('location', '')
            years_exp = user_info.get('years_exp', '2')
            
            # Top skills for context
            top_skills = [s.strip() for s in skills.split(',')[:4]] if skills else ['technology', 'problem-solving', 'collaboration']
            
            # Vary the prompt each time
            variation = random.choice([
                {
                    "angle": "business impact angle",
                    "emphasis": "measurable business outcomes and ROI"
                },
                {
                    "angle": "technical expertise angle", 
                    "emphasis": "architectural excellence and technical depth"
                },
                {
                    "angle": "innovation angle",
                    "emphasis": "forward-thinking solutions and emerging technologies"
                }
            ])
            
            prompt = f"""Generate an EXCEPTIONAL professional summary using the {variation["angle"]}.

CONTEXT:
- Position: {profession}
- Years: {years_exp} years
- Top Skills: {', '.join(top_skills)}
- Emphasis: {variation["emphasis"]}
- Variation Token: {variation_nonce}
- Recent Phrases To Avoid:
{format_avoid_phrases(summary_avoid_phrases)}

MUST CREATE 3-4 SENTENCES with unique phrasing:
1. Start with a strong adjective (different each time - Results-driven, Innovative, Proven, Strategic, Dynamic, etc.)
2. Include years and specific domain expertise
3. Highlight 2-3 key competencies showing {variation["emphasis"]}
4. Use strong action verbs (Spearheaded, Architected, Delivered, Engineered, Optimized)
5. Show concrete impact and business value
6. Unique ending that varies structure
7. Avoid wording used in Recent Phrases To Avoid

NO: Personal pronouns, generic phrases, weak verbs
OUTPUT ONLY THE SUMMARY - NO PREAMBLE."""
            
            try:
                print(f"[DEBUG] Generating professional summary ({variation['angle']})...")
                llm = get_llm_resume()
                if llm is None:
                    raise ValueError("LLM instance is not available")
                res = llm.invoke([HumanMessage(content=prompt)])
                content = res.content if isinstance(res.content, str) else str(res.content)
                result = content.strip().replace('```', '').replace('**', '')
                
                if result and len(result.strip()) > 30:
                    print(f"[SUCCESS] Generated summary: {result[:80]}...")
                    return result
                    
            except Exception as e:
                print(f"[WARN] Summary generation failed: {e}")
            
            # Vary fallback templates
            primary_skill = top_skills[0] if top_skills else "technology"
            fallback_templates = [
                f"Results-driven {profession} with {years_exp}+ years of proven expertise in delivering scalable solutions. Specialized in {primary_skill} with demonstrated ability to optimize systems and drive measurable business outcomes. Committed to continuous learning and leveraging innovative technologies to solve complex challenges.",
                
                f"Strategic {profession} with {years_exp}+ years of hands-on experience architecting enterprise-scale solutions. Deep expertise in {primary_skill} and proven track record of translating business requirements into technical excellence. Passionate about innovation and delivering sustainable competitive advantage.",
                
                f"Innovative {profession} with {years_exp}+ years driving technical transformation and organizational growth. Strong foundation in {primary_skill} combined with strategic mindset for solving complex problems. Committed to fostering collaboration and mentoring teams to achieve exceptional results."
            ]
            return pick_non_repeating(fallback_templates, 'summary')
        
        def generate_job_description(user_info):
            """Generate high-impact achievement bullets - VARIED each time."""
            from langchain_core.messages import HumanMessage
            import random
            
            job_title = user_info.get('jobTitle', 'Software Engineer')
            employer = user_info.get('employer', 'Company')
            skills = user_info.get('skills', '')
            profession = user_info.get('profession', 'Software Engineer')
            
            # Extract top technologies
            tech_list = [s.strip() for s in skills.split(',')[:3]] if skills else []
            tech_str = ' and '.join(tech_list) if tech_list else 'modern technologies'
            
            # Vary the focus each time
            focus_angle = random.choice([
                "technical depth and architectural excellence",
                "business metrics and ROI (cost savings, revenue)",
                "team leadership and organizational impact",
                "scale and performance optimization"
            ])
            
            prompt = f"""Generate 4 EXCEPTIONAL achievement bullets with focus on {focus_angle}.

CONTEXT:
- Job Title: {job_title}
- Company: {employer}
- Tech Stack: {tech_str}
- Focus: {focus_angle}
- Variation Token: {variation_nonce}
- Recent Phrases To Avoid:
{format_avoid_phrases(experience_avoid_phrases)}

EACH BULLET MUST:
1. Start with DIFFERENT power verb: Engineered, Architected, Led, Optimized, Delivered, Spearheaded, Implemented, Designed
2. Include SPECIFIC metrics: percentages, cost savings, time reductions, scale metrics, or quantifiable improvements
3. Include technology name
4. Different metric types across bullets (percentage, money, time, and scale)
5. One line, professional, no personal pronouns
6. Avoid wording used in Recent Phrases To Avoid

STRUCTURE VARIATIONS:
- Verb + detail using tech, metric result
- Verb + achievement, metric plus context
- Verb + team or scope, resulting in metric
- Verb + initiative, reducing or improving metric

OUTPUT - 4 varied bullets starting with dash, NO PREAMBLE."""
            
            try:
                print(f"[DEBUG] Generating job description ({focus_angle})...")
                llm = get_llm_resume()
                if llm is None:
                    raise ValueError("LLM instance is not available")
                res = llm.invoke([HumanMessage(content=prompt)])
                content = res.content if isinstance(res.content, str) else str(res.content)
                result = content.strip()
                
                # Ensure we have bullet format
                if '-' not in result:
                    result = '\n'.join(f"- {line.strip()}" for line in result.split('\n') if line.strip())
                
                if result and len(result.strip()) > 20:
                    print(f"[SUCCESS] Generated bullets: {result[:80]}...")
                    return result
                    
            except Exception as e:
                print(f"[WARN] Job description generation failed: {e}")
            
            # Vary fallback bullets
            tech = tech_list[0] if tech_list else 'modern technologies'
            fallback_options = [
                f"- Engineered scalable {tech} solutions improving system uptime from 95% to 99.9%\n" +
                "- Led migration project enabling 3x user capacity with 40% infrastructure cost reduction\n" +
                "- Optimized API response times by 55%, enhancing user experience and retention rates\n" +
                "- Implemented automated testing reducing production bugs by 65% and deployment time by 75%",
                
                f"- Architected microservices using {tech} processing 5M+ daily transactions with under 100ms latency\n" +
                "- Spearheaded performance optimization initiative reducing cloud costs by 60K annually\n" +
                "- Led team of 4 engineers delivering 8 features on schedule improving customer satisfaction by 35%\n" +
                "- Designed CI/CD pipeline enabling 50+ deployments weekly with zero downtime releases",
                
                f"- Delivered {tech} platform handling 50K concurrent users improving application stability by 45%\n" +
                "- Optimized database queries and caching strategy reducing average load time from 4s to 0.8s\n" +
                "- Led architecture redesign supporting 10x business growth and expanding market opportunities\n" +
                "- Implemented comprehensive monitoring and alerting reducing incident resolution time by 70%",
                
                f"- Engineered high-throughput {tech} system processing 100K+ records daily with 99.99% reliability\n" +
                "- Led cross-team initiative reducing operational overhead by 35% and improving code maintainability\n" +
                "- Architected event-driven architecture enabling real-time analytics reducing latency by 80%\n" +
                "- Mentored 2 junior developers improving code quality metrics by 40% and team productivity"
            ]
            return pick_non_repeating(fallback_options, 'experience')

        # Prepare user context for AI enhancement
        user_context = {
            'profession': data.get('profession', 'Software Engineer'),
            'skills': data.get('skills', ''),
            'jobTitle': data.get('jobTitle', ''),
            'employer': data.get('employer', ''),
            'location': data.get('location', '')
        }

        summary_text = enhance_text(data.get('summary', ''), "professional summary", user_context)
        job_desc_text = enhance_text(data.get('jobDesc', ''), "work experience description/achievements", user_context)

        # 2. PROFESSIONAL SUMMARY
        if summary_text:
            add_section_title("Professional Summary")
            pdf.set_font('Arial', '', 11)
            pdf.multi_cell(0, 6, sanitize_for_fpdf(summary_text))
            
        # 3. EXPERIENCE
        if data.get('jobTitle') or data.get('employer'):
            add_section_title("Work Experience")
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(100, 6, sanitize_for_fpdf(data.get('jobTitle', 'Role')), ln=0)
            
            # Dates aligned right
            pdf.set_font('Arial', 'I', 11)
            dates = f"{data.get('jobStart', '')} - {data.get('jobEnd', '')}"
            pdf.cell(0, 6, sanitize_for_fpdf(dates), ln=True, align='R')
            
            # Employer
            pdf.set_font('Arial', 'I', 11)
            pdf.cell(0, 6, sanitize_for_fpdf(data.get('employer', 'Company')), ln=True)
            
            # Description
            pdf.set_font('Arial', '', 11)
            pdf.multi_cell(0, 6, sanitize_for_fpdf(job_desc_text))

        # 4. EDUCATION
        if data.get('school') or data.get('degree'):
            add_section_title("Education")
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(100, 6, sanitize_for_fpdf(data.get('school', 'University')), ln=0)
            
            pdf.set_font('Arial', 'I', 11)
            pdf.cell(0, 6, sanitize_for_fpdf(data.get('gradYear', '')), ln=True, align='R')
            
            pdf.set_font('Arial', '', 11)
            deg_text = f"{data.get('degree', '')} - {data.get('schoolLoc', '')}"
            pdf.cell(0, 6, sanitize_for_fpdf(deg_text), ln=True)

        # 5. SKILLS
        if data.get('skills'):
            add_section_title("Skills & Competencies")
            pdf.set_font('Arial', '', 11)
            pdf.multi_cell(0, 6, sanitize_for_fpdf(data.get('skills', '')))
            
        session_id = str(uuid.uuid4())
        filepath = os.path.join(UPLOAD_FOLDER, f"{session_id}.pdf")

        _remember_resume_variation(user.get('_id'), summary_text, job_desc_text)

        # Extract skills for DB
        from mock_interview import extract_skills
        combined_text = f"{data.get('profession', '')} {data.get('summary', '')} {data.get('jobDesc', '')} {data.get('degree', '')} {data.get('skills', '')}"
        extracted = extract_skills(combined_text)

        # ──────────────────────────────────────────────
        # ATS Score & LaTeX Generation
        # ──────────────────────────────────────────────
        def calc_ats_score():
            s = 20 # Base
            if len(summary_text or "") > 150: s += 15
            if len(job_desc_text or "") > 250: s += 20
            if len(extracted) > 8: s += 20
            if data.get('degree'): s += 10
            if data.get('email') and data.get('phone'): s += 15
            return min(98, s)

        def sanitize_latex(text):
            if not text: return ""
            return text.replace('&', '\\&').replace('%', '\\%').replace('$', '\\$').replace('#', '\\#')

        def get_latex():
            summary = sanitize_latex(summary_text)
            experience = sanitize_latex(job_desc_text)
            skills = sanitize_latex(data.get('skills', ''))
            
            return r"""\documentclass[11pt,a4paper,sans]{{moderncv}}
\moderncvstyle{{classic}}
\moderncvcolor{{blue}}
\usepackage[utf8]{{inputenc}}
\usepackage[scale=0.75]{{geometry}}

\name{{{firstName}}}{{{lastName}}}
\title{{{profession}}}
\address{{{location}}}
\phone[mobile]{{{phone}}}
\email{{{email}}}

\begin{{document}}
\makecvtitle

\section{{Summary}}
{summary}

\section{{Experience}}
\cventry{{{jobStart}--{jobEnd}}}{{{jobTitle}}}{{{employer}}}{{}}{{}}{{{experience}}}

\section{{Education}}
\cventry{{{gradYear}}}{{{degree}}}{{{school}}}{{{schoolLoc}}}{{}}{{}}

\section{{Skills}}
\cvitem{{Technical}}{{{skills}}}

\end{{document}}""".format(
                firstName=data.get('firstName', ''),
                lastName=data.get('lastName', ''),
                profession=data.get('profession', 'Software Engineer'),
                location=data.get('location', 'Global'),
                phone=data.get('phone', ''),
                email=data.get('email', ''),
                summary=summary,
                jobStart=data.get('jobStart', ''),
                jobEnd=data.get('jobEnd', 'Present'),
                jobTitle=data.get('jobTitle', 'Role'),
                employer=data.get('employer', 'Company'),
                experience=experience,
                gradYear=data.get('gradYear', ''),
                degree=data.get('degree', ''),
                school=data.get('school', ''),
                schoolLoc=data.get('schoolLoc', ''),
                skills=skills
            )

        ats_score = calc_ats_score()
        latex_code = get_latex()

        import requests
        pdf_compiled = False
        try:
            print("[INFO] Attempting to compile LaTeX via texlive.net...")
            files = {'filecontents[]': (None, latex_code), 'filename[]': (None, 'document.tex')}
            api_data = {'engine': 'pdflatex', 'return': 'pdf'}
            res = requests.post("https://texlive.net/cgi-bin/latexcgi", data=api_data, files=files, timeout=25)
            if res.status_code == 200 and b'%PDF' in res.content[:10]:
                with open(filepath, "wb") as f:
                    f.write(res.content)
                pdf_compiled = True
                print("[SUCCESS] LaTeX compiled and saved successfully!")
            else:
                print(f"[WARN] LaTeX API failed or returned non-PDF. Status: {res.status_code}")
        except Exception as e:
            print("[WARN] LaTeX API request error:", e)

        if not pdf_compiled:
            print("[INFO] Falling back to basic FPDF generation.")
            pdf.output(filepath)

        # Update user session with the new resume and score
        db.update_user(user['_id'], 
                       session_id=session_id, 
                       extracted_skills=extracted)
        
        return jsonify({
            'status': 'success', 
            'session_id': session_id, 
            'skills': extracted, 
            'ats_score': ats_score,
            'latex': latex_code
        })
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/download-resume', methods=['GET'])
def download_resume():
    try:
        token = request.args.get('token')
        if not token:
            return "Unauthorized", 401
        user = db.get_user_by_token(token)
        if not user or not user.get('session_id'):
            return "No resume found", 404
        
        session_id = user.get('session_id')
        filepath = os.path.join(UPLOAD_FOLDER, f"{session_id}.pdf")
        if os.path.exists(filepath):
            from flask import send_file
            return send_file(filepath, as_attachment=True, download_name="My_AI_Resume.pdf", mimetype='application/pdf')
        return "File not found", 404
    except Exception as e:
        return str(e), 500

@app.route('/api/dashboard/analyze-resume', methods=['POST'])
def analyze_resume_for_dashboard():
    """
    Full pipeline:
      1. Extract text from PDF using pdfminer.six / PyPDF2 / pytesseract OCR
      2. Extract skills from text using the curated SKILLS_DB
      3. Scrape real internship listings from the internet for each top skill
      4. Score each listing against user's skill profile
      5. Return stats + top matched internships
    """
    user_email = None
    user_name = None
    from mock_interview import extract_text_from_pdf, extract_skills, infer_role, extract_resume_profile, get_skill_gap_recommendations
    import requests
    from bs4 import BeautifulSoup

    stream = get_resume_stream_from_req(request)
    if not stream:
        return jsonify({"error": "No resume found. Please provide session_id or file."}), 400

    try:
        # ── Step 1: Extract text ──────────────────────────────────────
        try:
            text = extract_text_from_pdf(stream)
        except Exception as e:
            return jsonify({"error": f"Text extraction failed: {str(e)}"}), 500

        if not text or len(text.strip()) < 30:
            # Fallback: Try to use existing profile data or infer from user metadata if extraction fails
            auth_header = request.headers.get('Authorization', '')
            if auth_header.startswith('Bearer '):
                token = auth_header.replace('Bearer ', '').strip()
                user = db.get_user_by_token(token)
                if user:
                    saved_skills = user.get('extractedSkills', [])
                    if not saved_skills:
                        # If no skills saved, infer from target_role (case-insensitive)
                        from mock_interview import ROLE_MAP
                        target = user.get('target_role', 'Software Engineer')
                        # Try exact match, then try title-case
                        saved_skills = ROLE_MAP.get(target) or ROLE_MAP.get(target.title()) or ['Python', 'Communication', 'Problem Solving']
                    
                    text = f"Profile for {user.get('full_name') or 'User'}. Target Role: {user.get('target_role') or 'Software Engineer'}. Skills: {', '.join(saved_skills)}"
                else:
                    return jsonify({"error": "Could not extract readable text from this PDF and no saved profile found. Please try a text-based PDF."}), 400
            else:
                return jsonify({"error": "Could not extract readable text from this PDF. Please try a text-based PDF."}), 400

        # ── Step 2: Extract skills, role, and profile details ────────
        profile = extract_resume_profile(text)
        skills = profile.get("skills", [])
        if not isinstance(skills, list):
            skills = []
        inferred_role = profile.get("role") or "Software Engineer"

        # DEEP INFERENCE: If skills are not directly mentioned, use AI to infer from experience/projects
        if len(skills) < 5:
            print("[Analysis] Basic extraction weak. Running Deep AI Inference...")
            llm = get_llm()
            if llm:
                try:
                    from langchain_core.messages import HumanMessage
                    inference_prompt = f"""
                    Analyze this resume text and identify the core technical skills and professional role.
                    Even if there isn't a 'Skills' section, infer them from the Projects, Experience, and Education descriptions.
                    
                    RESUME TEXT:
                    {text[:3000]}
                    
                    Return ONLY a JSON object:
                    {{
                        "inferred_skills": ["skill1", "skill2", ...],
                        "suggested_role": "Professional Role Title"
                    }}
                    """
                    resp = llm.invoke([HumanMessage(content=inference_prompt)])
                    content = resp.content if isinstance(resp.content, str) else str(resp.content)
                    raw_resp = content.strip()
                    if "```json" in raw_resp: raw_resp = raw_resp.split("```json")[1].split("```")[0].strip()
                    elif "```" in raw_resp: raw_resp = raw_resp.split("```")[1].split("```")[0].strip()
                    
                    deep_data = json.loads(raw_resp)
                    ai_skills = deep_data.get("inferred_skills", [])
                    if ai_skills:
                        skills = list(set(skills + ai_skills))
                    if deep_data.get("suggested_role"):
                        inferred_role = deep_data["suggested_role"]
                    print(f"[Analysis] Deep Inference found: {skills}")
                except Exception as e:
                    print(f"[Analysis] Deep Inference failed: {e}")

        if not skills:
            skills = ['Python', 'Communication', 'Problem Solving']

        # Update user profile and save file if provided
        update_data = None
        user = None
        auth_header = request.headers.get('Authorization', '')
        if auth_header.startswith('Bearer '):
            token = auth_header.replace('Bearer ', '').strip()
            user = db.get_user_by_token(token)
        
        if not user:
            # Fallback to session_id lookup if no Bearer token
            sid = request.form.get('session_id') or (request.get_json(silent=True) or {}).get('session_id')
            if sid:
                user = db.get_user_by_session_id(sid)

        if user:
            update_data = {
                'extracted_skills': skills,
                'extracted_role': inferred_role
            }
            
            # If a new resume was uploaded, save it and update session_id
            if 'resume' in request.files and request.files['resume'].filename != '':
                import uuid
                new_session_id = str(uuid.uuid4())
                filepath = os.path.join(UPLOAD_FOLDER, f"{new_session_id}.pdf")
                # Seek to beginning to ensure we save the whole file (in case extraction already read it)
                request.files['resume'].seek(0)
                request.files['resume'].save(filepath)
                update_data['session_id'] = new_session_id
            
            db.update_user(user['_id'], **update_data)
            
            # Capture for email
            user_name = user.get('full_name') or user.get('fullName') or "User"

        # ── Step 2.5: Skill Gap Analysis ─────────────────────────────
        skill_gap_data = get_skill_gap_recommendations(skills, inferred_role)

        # ── Step 3: Scrape real internships ──────────────────────────
        def scrape_internships(skill_query, role_query, fast_mode=False):
            """Scrape real internship listings. fast_mode uses only 1 source."""
            from tools import search_internships
            import json
            
            query = f"{role_query} internship in India knowing {skill_query}"
            
            # If in fast_mode (for bot/quick views), we limit sources to just LinkedIn
            # or a very fast search to keep latency low.
            if fast_mode:
                query += " site:linkedin.com"
            
            print(f"[Dashboard Scraper] {'FAST ' if fast_mode else ''}Search for: {query}")
            
            try:
                # Use the tool with restricted sources for fast_mode
                tool_output = getattr(search_internships, "invoke")({"query": query})
                
                # Extract JSON from markdown response
                json_str = tool_output.replace("```internship_cards\n", "").replace("\n```", "").strip()
                cards = json.loads(json_str)
                
                formatted = []
                for card in cards:
                    formatted.append({
                        'title': card.get('title', 'Internship Role'),
                        'company': card.get('company', 'Tech Company'),
                        'location': card.get('location', 'India'),
                        'snippet': card.get('description', ''),
                        'url': card.get('apply_link', '#')
                    })
                return formatted
            except Exception as e:
                print(f"[Dashboard Scraper] Tool error: {e}")
                return []

        # ── Step 4: Score each scraped job against skills ─────────────
        def compute_match_score(job_text, user_skills_lower):
            """Returns a match percentage (0-100) based on skill overlap."""
            job_lower = job_text.lower()
            if not user_skills_lower: return 50 # Neutral if no skills
            matched = sum(1 for skill in user_skills_lower if skill.lower() in job_lower)
            base_score = min(100, 40 + int((matched / max(len(user_skills_lower), 1)) * 60))
            return base_score

        # Prepare queries
        is_fast = request.args.get('fast') == 'true' or request.form.get('fast') == 'true'
        skills_subset = skills[:3]
        skill_query = ", ".join(skills_subset)
        raw_jobs = scrape_internships(skill_query, inferred_role, fast_mode=is_fast)

        user_skills_lower = [s.lower() for s in skills]
        scored_jobs = []
        seen_urls = set()
        
        for job in raw_jobs:
            if job['url'] in seen_urls: continue
            seen_urls.add(job['url'])
            
            combined_text = job['title'] + ' ' + job['snippet']
            score = compute_match_score(combined_text, user_skills_lower)

            # Bonus for role match
            role_parts = str(inferred_role).lower().split()
            safe_title = str(job.get('title', '')).lower()
            if role_parts and role_parts[0] in safe_title:
                score = min(99, score + 5)

            scored_jobs.append({
                'title': job['title'][:70],
                'snippet': job['snippet'][:150],
                'url': job['url'],
                'matchScore': score,
                'company': job['company'],
                'mode': 'Remote' if 'Remote' in job.get('location', '') else 'On-site',
                'location': job.get('location', 'India')
            })

        scored_jobs.sort(key=lambda j: j['matchScore'], reverse=True)
        top_jobs = scored_jobs[:10]

        # Fix missing all_job_lower for market trend analysis
        all_job_lower = " ".join([j['title'] + ' ' + j['snippet'] for j in scored_jobs]).lower()

        # Step 5: Compute dashboard stats
        matched_count = len(scored_jobs)
        avg_score = round(sum(j['matchScore'] for j in top_jobs) / max(len(top_jobs), 1)) if top_jobs else 75
        interviews_completed = max(1, min(20, len(skills) // 3))

        # Step 6: Skill Gap Analysis with Resources
        gap_results = get_skill_gap_recommendations(skills, inferred_role)
        gap_skills = []
        for rec in gap_results["recommendations"]:
            gap_skills.append({
                'skill': rec['skill'],
                'priority': rec.get('priority', 'high'),
                'demand': rec.get('demand', 3),
                'youtube': rec['youtube'],
                'books': rec['books']
            })

        # Step 7: Live Market Trends
        ROLE_MAP = {
            'Software Engineer': ['software engineer','backend','frontend','full stack','fullstack'],
            'Data Science':      ['data science','data scientist','data analyst','analytics'],
            'Machine Learning':  ['machine learning','ml engineer','ai engineer','deep learning'],
            'Web Development':   ['web developer','react developer','frontend developer'],
            'Mobile Dev':        ['mobile','android','ios','flutter','react native'],
            'DevOps / Cloud':    ['devops','cloud engineer','kubernetes','aws','azure'],
            'UI/UX Design':      ['ui/ux','ui design','ux design','product design','figma'],
            'Cybersecurity':     ['cybersecurity','security analyst','penetration','ethical hacking'],
        }
        role_counts = {r: sum(all_job_lower.count(kw) for kw in kws) for r, kws in ROLE_MAP.items()}
        total_hits = max(sum(role_counts.values()), 1)
        market_trends = sorted(
            [{'role': r, 'count': c, 'percent': round((c / total_hits) * 100)}
             for r, c in role_counts.items() if c > 0],
            key=lambda x: x['count'], reverse=True
        )[:8]

        # Step 8: Company Culture Insights
        def _company_insight(snippet, company):
            s = (snippet + ' ' + company).lower()
            if any(w in s for w in ['google','microsoft','amazon','meta','apple','netflix','openai']):
                return {'tag': 'FAANG-tier',  'color': '#f59e0b', 'desc': 'Top-tier tech - rigorous DSA rounds, competitive pay'}
            if any(w in s for w in ['startup','seed','series a','early stage','venture']):
                return {'tag': 'Startup',     'color': '#a855f7', 'desc': 'Startup culture - high ownership, fast-paced, direct impact'}
            if any(w in s for w in ['remote','work from home','wfh','distributed']):
                return {'tag': 'Remote-First','color': '#00f0ff', 'desc': 'Remote-first team - async culture, flexible hours'}
            if any(w in s for w in ['agile','scrum','sprint','kanban']):
                return {'tag': 'Agile',       'color': '#00e4b8', 'desc': 'Agile shop - sprints, standups, iterative delivery'}
            if any(w in s for w in ['research','phd','academic','publication']):
                return {'tag': 'Research',    'color': '#4da0ff', 'desc': 'Research-oriented - cutting-edge work, papers and experiments'}
            if any(w in s for w in ['fintech','finance','bank','trading','insurance']):
                return {'tag': 'FinTech',     'color': '#f97316', 'desc': 'Finance sector - compliance-aware, high-performance systems'}
            return     {'tag': 'Tech Co.',   'color': '#6366f1', 'desc': 'Tech company - collaborative culture, learning opportunities'}

        for job in top_jobs:
            job['insight'] = _company_insight(job.get('snippet', ''), job.get('company', ''))

        # --- Fallback: static market trends when scraping returns nothing ---
        if not market_trends:
            market_trends = [
                {'role': 'Software Engineer', 'count': 45, 'percent': 35},
                {'role': 'Data Science',      'count': 26, 'percent': 20},
                {'role': 'Machine Learning',  'count': 20, 'percent': 16},
                {'role': 'Web Development',   'count': 18, 'percent': 14},
                {'role': 'Mobile Dev',        'count': 10, 'percent': 8},
                {'role': 'DevOps / Cloud',    'count': 9,  'percent': 7},
            ]
        # --- Fallback: static matches when scraping returns nothing ---
        if not top_jobs:
            top_jobs = [
                {
                    'title': 'Software Engineering Intern',
                    'company': 'TechNova Solutions',
                    'location': 'Remote',
                    'mode': 'Remote',
                    'snippet': 'We are looking for a Software Engineering Intern with experience in Python and JavaScript to build scalable microservices. Must be familiar with Agile methodologies.',
                    'url': '#',
                    'matchScore': 92,
                    'insight': {'tag': 'Startup', 'color': '#a855f7', 'desc': 'Startup culture - high ownership, fast-paced, direct impact'}
                },
                {
                    'title': 'Data Science Intern',
                    'company': 'Global Analytics',
                    'location': 'Bengaluru, India',
                    'mode': 'Hybrid',
                    'snippet': 'Join our analytics team! You will work with Pandas, Scikit-Learn, and SQL to extract insights from massive datasets.',
                    'url': '#',
                    'matchScore': 88,
                    'insight': {'tag': 'Tech Co.', 'color': '#6366f1', 'desc': 'Tech company - collaborative culture, learning opportunities'}
                },
                {
                    'title': 'Full Stack Developer Intern',
                    'company': 'Innovate Inc.',
                    'location': 'Pune, India',
                    'mode': 'On-site',
                    'snippet': 'Seeking a talented Full Stack Intern. You will work with React, Node.js, and MongoDB. Experience with cloud platforms (AWS) is a plus.',
                    'url': '#',
                    'matchScore': 85,
                    'insight': {'tag': 'Agile', 'color': '#00e4b8', 'desc': 'Agile shop - sprints, standups, iterative delivery'}
                },
                {
                    'title': 'Cloud DevOps Intern',
                    'company': 'CloudOps Systems',
                    'location': 'Remote',
                    'mode': 'Remote',
                    'snippet': 'Learn cloud infrastructure automation using Kubernetes, Docker, and CI/CD pipelines. Remote-first team.',
                    'url': '#',
                    'matchScore': 81,
                    'insight': {'tag': 'Remote-First', 'color': '#00f0ff', 'desc': 'Remote-first team - async culture, flexible hours'}
                }
            ]
            for job in top_jobs:
                if 'insight' not in job:
                    job['insight'] = _company_insight(job.get('snippet', ''), job.get('company', ''))

        # Send email notification in background
        if user_email:
            import threading
            threading.Thread(target=send_analysis_email, args=(user_email, user_name, skills, inferred_role)).start()

        edu_data = profile.get("education")
        college_name = edu_data[0] if isinstance(edu_data, list) and edu_data else ""

        return jsonify({
            "name": profile.get("fullName", ""),
            "email": profile.get("email", ""),
            "phone": profile.get("phone", ""),
            "location": profile.get("location", ""),
            "experience_years": profile.get("yearsExperience", ""),
            "linkedin": profile.get("linkedin", ""),
            "github": profile.get("github", ""),
            "portfolio": profile.get("portfolio", ""),
            "bio": profile.get("summary", ""),
            "college": college_name,
            "skills": skills,
            "role": inferred_role,
            "stats": {
                "matchedInternships": matched_count,
                "interviewsCompleted": interviews_completed,
                "avgMatchScore": avg_score,
                "dayStreak": max(1, len(skills) // 4)
            },
            "topMatches": top_jobs,
            "skillGap": gap_skills,
            "marketTrends": market_trends,
            "session_id": (update_data.get('session_id') if update_data else request.form.get("session_id")),
            "rawTextPreview": text[:500] + ("..." if len(text) > 500 else "")
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Backend crash during analysis: {str(e)}"}), 500


def _extract_company(url, snippet):
    """Try to extract company name from URL or snippet."""
    if 'linkedin.com' in url:
        return 'LinkedIn Listing'
    if 'indeed.com' in url:
        return 'Indeed Listing'
    if 'internshala.com' in url:
        return 'Internshala'
    if 'glassdoor.com' in url:
        return 'Glassdoor'
    # Try to pull from snippet
    for word in ['at ', 'with ', '@ ']:
        if word in snippet.lower():
            idx = snippet.lower().index(word) + len(word)
            return snippet[idx:idx+20].split()[0].strip(',.') if idx < len(snippet) else 'Company'
    return 'Tech Company'


def _guess_mode(snippet):
    s = snippet.lower()
    if 'remote' in s:
        return 'Remote'
    if 'hybrid' in s:
        return 'Hybrid'
    return 'On-site'


@app.route("/api/generate-cover-letter", methods=["POST"])
def generate_cover_letter():
    try:
        payload = request.get_json(silent=True) if request.is_json else {}
        if not isinstance(payload, dict):
            payload = {}

        job_title = (
            request.form.get("jobTitle", "").strip()
            or request.form.get("job_title", "").strip()
            or str(payload.get("jobTitle", "")).strip()
            or str(payload.get("job_title", "")).strip()
            or "Software Engineer"
        )
        company = (
            request.form.get("company", "").strip()
            or str(payload.get("company", "")).strip()
            or "the Company"
        )
        job_description = (
            request.form.get("jobDescription", "").strip()
            or request.form.get("job_description", "").strip()
            or str(payload.get("jobDescription", "")).strip()
            or str(payload.get("job_description", "")).strip()
            or ""
        )
        from datetime import date
        today = date.today().strftime("%B %d, %Y")

        stream = get_resume_stream_from_req(request)
        if not stream:
            return jsonify({"error": "No resume found."}), 400

        # If user uploads a resume in this request, persist it as their latest resume
        # so subsequent cover-letter generations can reuse it automatically.
        uploaded_resume = request.files.get("resume")
        if uploaded_resume and uploaded_resume.filename:
            user = _get_user_from_request()
            if user:
                try:
                    latest_session_id = str(uuid.uuid4())
                    latest_path = os.path.join(UPLOAD_FOLDER, f"{latest_session_id}.pdf")
                    uploaded_resume.seek(0)
                    uploaded_resume.save(latest_path)
                    db.update_user(user["_id"], session_id=latest_session_id)
                except Exception as ex:
                    print(f"[Cover Letter] Could not persist uploaded resume: {ex}")

        from mock_interview import extract_text_from_pdf
        try:
            resume_text = extract_text_from_pdf(stream)
        except Exception as e:
            return jsonify({"error": f"Failed to read resume: {str(e)}"}), 500

        if len(resume_text.strip()) < 30:
            return jsonify({"error": "Could not extract readable text from this file."}), 400

        # Build strict official-format prompt with focus on uniqueness and professionalism
        prompt_text = f"""You are an elite executive career coach and professional cover letter writer. Generate a highly realistic, unique, and professional business cover letter.

Requirements:
- Use the exact date: {today}
- Target position: {job_title}
- Target company: {company}
- Tone: Confident, modern, engaging, and highly professional. Avoid generic clichés, overly formal archaic language, or robotic phrasing.
- Strategy: Do NOT just summarize the resume. Instead, tell a compelling brief story about the candidate's unique value proposition. Focus on impact, problem-solving, and how their specific background perfectly aligns with the needs of {company}.
- Use the candidate's details extracted from the resume below.

Resume Content:
---
{resume_text[:2500]}
---

Job Description (if provided):
---
{job_description[:1500] or "Not provided"}
---

The cover letter MUST follow this EXACT official format. YOU MUST REPLACE the bracketed placeholders below with the actual information extracted from the candidate's resume:

[Extract Candidate Full Name]
[Extract Candidate Email]
[Extract Candidate Phone, or omit line if not found]
[Extract City, or omit line if not found]
{today}

Hiring Manager
{company}
Subject: Application for the Position of {job_title}

Dear Hiring Manager,

[Opening paragraph: A strong, unique hook that expresses genuine enthusiasm for {company} and the {job_title} role. Avoid generic openings like "I am writing to apply for..."]

[Body paragraph(s): Highlight 2-3 specific, high-impact achievements from the resume that demonstrate immediate value to {company}. Connect the dots between past success and future potential. Use concrete metrics or outcomes where available.]

[Closing paragraph: A confident, forward-looking close. Reiterate interest, express eagerness to discuss how their skills align with {company}'s goals, and thank the reader.]

Sincerely,
[Candidate Full Name]

Return ONLY the formatted cover letter text, no extra explanations or markdown blocks around the text."""

        try:
            # Call configured LLM
            llm = get_llm()
            if llm is None:
                raise ValueError("LLM instance is not available")
            response = llm.invoke([HumanMessage(content=prompt_text)])
            cl_text = _normalize_llm_content(getattr(response, "content", ""))
            if not cl_text:
                raise ValueError("LLM returned empty cover letter content")
        except Exception as llm_error:
            print(f"[Cover Letter] LLM failed ({llm_error}), using fallback generator.")
            from mock_interview import extract_resume_profile
            profile = extract_resume_profile(resume_text)
            
            cand_name = profile.get("fullName") or "[Candidate Full Name]"
            cand_email = profile.get("email") or "[Candidate Email]"
            cand_phone = profile.get("phone") or "[Candidate Phone]"
            cand_location = profile.get("location") or "[Candidate Location]"
            skills = profile.get("skills", [])
            if not isinstance(skills, list):
                skills = []
            skills_str = ", ".join(str(s) for s in skills[:3]) if skills else "relevant technical skills"
            
            cl_text = f"{cand_name}\n{cand_email}\n{cand_phone}\n{cand_location}\n{today}\n\nHiring Manager\n{company}\nSubject: Application for the Position of {job_title}\n\nDear Hiring Manager,\n\nI am writing to express my strong enthusiasm for the {job_title} position at {company}. With my background and strong foundation in {skills_str}, I am confident in my ability to contribute effectively to your team and make an immediate impact.\n\nThroughout my academic and professional journey, I have developed a skill set that aligns closely with the requirements for this role. I have consistently demonstrated a commitment to delivering high-quality work and quickly adapting to new technical challenges. My past projects have taught me the importance of problem-solving, efficient execution, and continuous learning.\n\nI am particularly drawn to {company} because of your innovative approach and industry leadership. I would welcome the opportunity to discuss how my background and skills will be beneficial to your organization. Thank you for considering my application. I look forward to the possibility of an interview.\n\nSincerely,\n{cand_name}"

        # Return plain text for editing (use cover_letter key as expected by frontend)
        return jsonify({"cover_letter": cl_text, "text": cl_text})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Failed to generate cover letter: {str(e)}"}), 500


@app.route("/api/download-cover-letter", methods=["POST"])
def download_cover_letter():
    """Generate PDF or DOCX from user-edited cover letter text."""
    try:
        data = request.json
        text = data.get("text", "").strip()
        fmt = data.get("format", "pdf").lower()

        if not text:
            return jsonify({"error": "No text provided"}), 400

        if fmt == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("helvetica", size=11)

            # Encode the whole text safely to latin-1, then write in one shot
            safe_text = text.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 6, safe_text)

            pdf_bytes = bytes(pdf.output())
            from flask import send_file
            return send_file(io.BytesIO(pdf_bytes), mimetype='application/pdf', as_attachment=True, download_name='Cover_Letter.pdf')

        elif fmt == "docx":
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = docx_lib.Document()

            # Set margins
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.2)
                section.right_margin = Inches(1.2)

            lines = text.split('\n')
            for i, line in enumerate(lines):
                if not line.strip():
                    doc.add_paragraph()
                elif i < 6:  # Header block — bold
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip())
                    run.bold = True
                    run.font.size = Pt(11)
                else:
                    p = doc.add_paragraph(line.strip())
                    p.runs[0].font.size = Pt(11) if p.runs else None

            docx_io = io.BytesIO()
            doc.save(docx_io)
            docx_io.seek(0)
            from flask import send_file
            return send_file(docx_io, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name='Cover_Letter.docx')

        else:
            return jsonify({"error": "Invalid format. Use 'pdf' or 'docx'"}), 400

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Download generation failed: {str(e)}"}), 500




@app.route("/api/help-bot", methods=["POST"])
def help_bot():
    """AI assistant specialized in answering questions about the project sections."""
    try:
        data = request.get_json(silent=True) or {}
        user_message = data.get("message", "").strip()
        if not user_message:
            return jsonify({"error": "Empty message"}), 400

        # --- Personalization ---
        user = _get_user_from_request()
        user_name = user.get('full_name', 'there') if user else 'there'
        first_name = user_name.split()[0] if user_name != 'there' else 'friend'

        # --- Rule-based frank & friendly fallback ---
        msg_lower = user_message.lower()
        
        # Social & Identity
        if any(w in msg_lower for w in ["your name", "who are you"]):
            return jsonify({"response": f"I'm your **InternAI Guide**, {first_name}! Think of me as your personal career coach. I'm built to help you navigate this platform and land that dream internship. 🚀", "status": "success"})
        if any(w in msg_lower for w in ["hello", "hi", "hey"]):
            return jsonify({"response": f"Hey {first_name}! 👋 Great to see you. Ready to crush some internship applications today? How can I help?", "status": "success"})
        if any(w in msg_lower for w in ["thank", "thanks"]):
            return jsonify({"response": f"No worries, {first_name}! Always happy to help. You've got this! Anything else on your mind?", "status": "success"})
        if "bye" in msg_lower or "goodbye" in msg_lower:
            return jsonify({"response": f"Catch you later, {first_name}! Go get 'em! I'll be right here if you need me again. 👋", "status": "success"})
        if "how are you" in msg_lower:
            return jsonify({"response": f"I'm doing great, {first_name}! Especially when I'm helping talented people like you find their next big opportunity. How's your search going?", "status": "success"})
            
        # Feature guidance (Frank style)
        if "cover letter" in msg_lower:
            return jsonify({"response": f"Look, {first_name}, writing cover letters is a pain. That's why I've got the **Cover Letter AI** for you. Just head to your **Dashboard**, find 'Smart Writing', and let me handle the heavy lifting. It'll tailor everything to your resume in seconds!", "status": "success"})
        if "scan" in msg_lower or "job" in msg_lower or "internship" in msg_lower:
            return jsonify({"response": f"Searching for jobs is exhausting, I get it. Check out the **'AI Job Scanner'** in the sidebar. I've programmed it to hunt through LinkedIn, Unstop, and others so you don't have to. Just upload your resume and let me find the gold for you!", "status": "success"})
        if "interview" in msg_lower or "mock" in msg_lower:
            return jsonify({"response": f"Interview nerves are real, {first_name}! But practice makes perfect. Use the **'Mock Interview'** tool in the sidebar. It's a full-on simulation with an AI proctor that will actually give you feedback on how to improve.", "status": "success"})
        if "resume" in msg_lower or "analyze" in msg_lower:
            return jsonify({"response": f"Your resume is your ticket in! Click **'Analyze Saved Resume'** on the Bot page. I'll rip through it, find your best skills, and show you exactly where you stand in the market right now.", "status": "success"})

        system_prompt = f"""You are the InternAI Guide Bot, a frank, friendly, and highly personalized career assistant. 
        You are talking to {user_name}. Use their name occasionally to keep it personal.
        
        Platform Sections:
        1. AI Job Scanner: Scans multiple platforms (LinkedIn, Unstop, etc.) for live roles.
        2. Mock Interview AI: Real-time proctored simulation for interview practice.
        3. Cover Letter AI: Tailors professional letters to specific job descriptions.
        4. Dashboard: Central hub for match scores, skill gaps, and market trends.
        
        Style: Be very frank, honest, and supportive. Use a conversational tone as if you're a mentor. Use emojis.
        """

        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=user_message)
        ]
        
        ai_response = ""
        llm_instance = get_llm()
        if llm_instance:
            try:
                response = llm_instance.invoke(messages)
                ai_response = _normalize_llm_content(getattr(response, "content", ""))
            except:
                pass
        
        if not ai_response:
            ai_response = f"I'm here for you, {first_name}! Whether you want to scan for new jobs, practice for an interview, or just chat about your career path, just let me know what's up."
        
        return jsonify({
            "response": ai_response,
            "status": "success"
        })
    except Exception as e:
        print(f"Error in /api/help-bot: {e}")
        return jsonify({"error": str(e)}), 500


# Routes consolidated above


@app.route("/scan", methods=["POST"])
def scan_jobs():
    """API endpoint for the 'Start AI Scan' button — returns real job results."""
    import random
    from tools import (scrape_linkedin, scrape_indeed,
                       scrape_freshershub, scrape_internshiphub, scrape_placementindia,
                       scrape_unstop, scrape_social_media)

    data = request.get_json() or {}
    query = data.get("query", "software intern India")
    if "india" not in query.lower():
        query = f"{query} India"

    all_cards = []
    seen_links = set()
    scrapers = [
        scrape_linkedin, scrape_indeed,
        scrape_freshershub, scrape_internshiphub, scrape_placementindia,
        scrape_unstop, scrape_social_media
    ]
    for func in scrapers:
        try:
            for card in func(query, days_ago=7):
                link = card.get("apply_link", "")
                if link and link not in seen_links:
                    seen_links.add(link)
                    all_cards.append(card)
        except Exception as e:
            print(f"[scan] scraper error: {e}")

    all_cards = filter_india_jobs(all_cards)

    # Sort by applicants ascending (lowest competition first)
    all_cards.sort(key=lambda x: x.get("applicants", 10000))
    for i, card in enumerate(all_cards):
        card["id"] = i + 1

    return jsonify({"jobs": all_cards})


def _normalize_llm_content(content):
    """Normalize provider-specific LLM content blocks into plain text."""
    if content is None:
        return ""
    if isinstance(content, str):
        return content.strip()
    if isinstance(content, dict):
        text = content.get("text") or content.get("content") or content.get("output_text") or ""
        return str(text).strip()
    if isinstance(content, list):
        parts = []
        for item in content:
            if isinstance(item, str):
                text = item.strip()
                if text:
                    parts.append(text)
                continue
            if isinstance(item, dict):
                text = item.get("text") or item.get("content") or item.get("output_text") or ""
                if text:
                    parts.append(str(text).strip())
                continue
            text = getattr(item, "text", "") or getattr(item, "content", "")
            if text:
                parts.append(str(text).strip())
        return "\n".join(p for p in parts if p).strip()
    return str(content).strip()


@app.route("/api/bot/sessions", methods=["GET"])
def get_sessions():
    user = _get_user_from_request()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    from database import get_user_chat_sessions
    sessions = get_user_chat_sessions(user["_id"])
    return jsonify({"sessions": sessions, "status": "success"})

@app.route("/api/bot/chat/<session_id>", methods=["GET"])
def get_session_chat(session_id):
    user = _get_user_from_request()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    from database import get_chat_messages
    messages = get_chat_messages(session_id)
    return jsonify({"messages": messages, "status": "success"})

@app.route("/api/bot/chat/<session_id>", methods=["DELETE"])
def delete_session_chat(session_id):
    user = _get_user_from_request()
    if not user:
        return jsonify({"error": "Unauthorized"}), 401
    from database import delete_chat_session
    delete_chat_session(session_id, user["_id"])
    return jsonify({"status": "success"})


@app.route("/chat", methods=["POST"])
def chat():
    """Handle chat messages from the user."""
    try:
        data = request.get_json(silent=True) or {}
        user_message = data.get("message", "").strip()
        session_id = data.get("session_id", "").strip()

        if not user_message:
            return jsonify({"error": "Empty message"}), 400

        user = _get_user_from_request()
        user_id = user["_id"] if user else ""
        
        if not session_id or session_id == "default":
            from database import create_chat_session
            title = user_message[:30] + "..." if len(user_message) > 30 else user_message
            session_id = create_chat_session(user_id, title)
            
        from database import save_chat_message
        save_chat_message(session_id, "user", user_message, user_id)

        # Get chat history
        get_chat_history(session_id)
        append_chat_messages(session_id, HumanMessage(content=user_message))
        history = get_chat_history_snapshot(session_id)

        # --- LIGHTNING-FAST TOOL INTERCEPTOR ---
        # Don't wait 15-30 seconds for the local CPU LLM to process if we already know they want jobs!
        search_keywords = ["intern", "job", "role", "developer", "engineer", "position", "work", "find", "show", "remote"]
        needs_search = any(k in user_message.lower() for k in search_keywords)
        
        ai_response = ""
        
        if needs_search:
            # ── SOURCE FILTER: detect if the user wants a specific platform only ──
            # Determine source filter
            msg_lower = user_message.lower()
            discovered_sources = []
            
            source_map = {
                "linkedin": "linkedin",
                "indeed": "indeed",
                "freshershub": "freshershub",
                "placementhub": "placementhub",
                "internshiphub": "internshiphub",
                "unstop": "unstop",
                "twitter": "social",
                "facebook": "social",
                "instagram": "social",
                "social": "social",
                "placementindia": "placementindia",
            }
            
            for kw, src in source_map.items():
                if kw in msg_lower:
                    discovered_sources.append(src)
            # Deduplicate (e.g. twitter + facebook both map to "social")
            discovered_sources = list(dict.fromkeys(discovered_sources))
                    
            sources_param = ",".join(discovered_sources) if discovered_sources else "linkedin,indeed,freshershub,placementhub,internshiphub,unstop,social"

            tool_output = getattr(search_internships, "invoke")({
                "query": user_message,
                "sources": sources_param
            })
            
            source_labels = discovered_sources if discovered_sources else ["All Sources"]
            source_title = "/".join(s.capitalize() for s in source_labels)
            
            if "```internship_cards" in tool_output:
                ai_response = f"Here are live jobs from **{source_title}**:\n\n" + tool_output
            else:
                ai_response = tool_output
        else:
            # Only run the LLM natively if it's conversational small talk
            llm = get_llm()
            if llm is None:
                raise ValueError("LLM instance is not available")
            llm_with_tools = llm.bind_tools(tools)
            ai_msg = llm_with_tools.invoke(history)
            
            if ai_msg.tool_calls:
                first_call = ai_msg.tool_calls[0] if ai_msg.tool_calls else {}
                if isinstance(first_call, dict):
                    call_args = first_call.get("args", {}) or {}
                else:
                    call_args = getattr(first_call, "args", {}) or {}
                query_arg = call_args.get("query", user_message) if isinstance(call_args, dict) else user_message
                tool_output = getattr(search_internships, "invoke")({"query": query_arg, "sources": "linkedin,indeed,freshershub,placementhub,internshiphub,unstop,social"})
                if "```internship_cards" in tool_output:
                    ai_response = "Here are the live job listings:\n\n" + tool_output
                else:
                    ai_response = tool_output
            else:
                ai_response = _normalize_llm_content(ai_msg.content)
                if ai_response.startswith("{") and ai_response.endswith("}"):
                    ai_response = "I couldn't quite understand that. Are you looking for a specific internship?"

        if not ai_response:
            ai_response = "I'm sorry, I couldn't process that. Could you try rephrasing? 😊"

        # Update history with AI's tool or text response
        if "```internship_cards" in ai_response:
            # Do NOT store 3,000 tokens of JSON in the LLM's chat history! It will freeze local models.
            append_chat_messages(session_id, AIMessage(content="I successfully found and displayed the matching internship cards to the user."))
            save_chat_message(session_id, "ai", ai_response, user_id)
        else:
            append_chat_messages(session_id, AIMessage(content=ai_response))
            save_chat_message(session_id, "ai", ai_response, user_id)

        return jsonify({
            "response": ai_response,
            "session_id": session_id,
            "status": "success"
        })

    except Exception as e:
        print(f"Error in /chat: {e}")
        error_msg = str(e)
        if "connection error" in error_msg.lower() or "connectionrefusederror" in error_msg.lower():
            return jsonify({
                "response": "⚠️ **API Error**: Could not connect to the AI model. Please check your internet connection or API key status.",
                "status": "success"
            })
        return jsonify({
            "error": f"Something went wrong: {str(e)}",
            "status": "error"
        }), 500


@app.route("/upload-resume", methods=["POST"])
def upload_resume():
    """
    Handle PDF resume uploads:
    1. Extract text via PyPDF2
    2. Parse key skills with a fast local regex extractor (NO LLM = instant)
    3. Search all platforms for jobs posted in the LAST 7 DAYS
    4. Return job cards + a skill summary
    """
    import re

    try:
        stream = get_resume_stream_from_req(request)
        if not stream:
            return jsonify({"error": "No resume found"}), 400

        session_id = request.form.get("session_id", "default")

        # ── Step 1: Extract text from PDF ─────────────────────────────────
        reader = PdfReader(stream)
        resume_text = ""
        for page in reader.pages:
            text = page.extract_text()
            if text:
                resume_text += text + "\n"

        if not resume_text.strip():
            return jsonify({"error": "Could not extract text from the PDF. Please try a text-based PDF."}), 400

        # ── Step 2: Fast keyword extractor (NO LLM — instant) ─────────────
        text_lower = resume_text.lower()

        TECH_SKILLS = [
            "python", "javascript", "typescript", "java", "c++", "c#", "go", "rust",
            "react", "angular", "vue", "node.js", "nodejs", "django", "flask", "fastapi",
            "spring", "express", "next.js", "nextjs", "sql", "mysql", "postgresql",
            "mongodb", "redis", "machine learning", "deep learning", "nlp",
            "computer vision", "data science", "tensorflow", "pytorch", "keras",
            "scikit-learn", "pandas", "numpy", "aws", "azure", "gcp", "docker",
            "kubernetes", "git", "html", "css", "rest api", "graphql",
            "android", "ios", "flutter", "react native", "swift", "kotlin",
            "blockchain", "solidity", "web3", "devops"
        ]

        ROLE_KEYWORDS = [
            "software engineer", "software developer", "backend developer",
            "frontend developer", "full stack developer", "data scientist",
            "data analyst", "ml engineer", "machine learning engineer",
            "ai engineer", "devops engineer", "android developer", "ios developer",
            "mobile developer", "web developer", "intern", "fresher"
        ]

        found_skills = [s for s in TECH_SKILLS if s in text_lower]
        found_roles = [r for r in ROLE_KEYWORDS if r in text_lower]

        # ── Step 3: Build ROTATING query variations ────────────────────────
        # Shuffle skills so every upload picks different combos as primary keys
        import random
        shuffled_skills = found_skills[:] if found_skills else ["software"]
        random.shuffle(shuffled_skills)

        primary_role = found_roles[0] if found_roles else "developer"

        # Build up to 3 unique queries from different skill combos
        queries = []
        if len(shuffled_skills) >= 2:
            queries.append(f"{shuffled_skills[0]} {primary_role} intern india")
            queries.append(f"{shuffled_skills[1]} developer intern india")
        if len(shuffled_skills) >= 3:
            queries.append(f"{shuffled_skills[2]} engineer internship india")
        if not queries:
            queries = [f"{shuffled_skills[0]} {primary_role} intern india"]

        search_query = queries[0]  # primary query for display

        # ── Step 4: Multi-query parallel scraping with STRICT 7-day filter ──
        from concurrent.futures import ThreadPoolExecutor, as_completed
        from tools import scrape_linkedin, scrape_indeed, \
                          scrape_freshershub, scrape_internshiphub, scrape_placementindia, \
                          scrape_unstop, scrape_social_media

        seen_links = set()
        all_cards = []

        scrapers = [
            scrape_linkedin, scrape_indeed, 
            scrape_freshershub, scrape_internshiphub, scrape_placementindia,
            scrape_unstop, scrape_social_media
        ]

        def run_scraper(s_func, q):
            try:
                return s_func(q, days_ago=7)
            except Exception as e:
                print(f"Scraper error for query '{q}': {e}")
                return []

        # Execute all query/scraper combinations in parallel
        with ThreadPoolExecutor(max_workers=15) as executor:
            futures = []
            for q in queries:
                for s_func in scrapers:
                    futures.append(executor.submit(run_scraper, s_func, q))
            
            for future in as_completed(futures):
                results = future.result()
                for card in results:
                    link = card.get("apply_link", "")
                    if link and link not in seen_links:
                        seen_links.add(link)
                        all_cards.append(card)

        all_cards = filter_india_jobs(all_cards)

        # Shuffle the final pool to ensure a different visual order every time
        random.shuffle(all_cards)

        # Re-number and badge (after shuffle)
        for i, card in enumerate(all_cards):
            card["id"] = i + 1
            source = card.pop("source", "Job Board")
            card["title"] = f"[{source}] {card['title']}"

        # ── Step 4: Build the response ─────────────────────────────────────
        skills_display = ", ".join(f"**{s.title()}**" for s in found_skills[:8]) or "General Software Skills"
        roles_display = ", ".join(r.title() for r in found_roles[:3]) or "Software Developer"

        if all_cards:
            cards_block = "```internship_cards\n" + json.dumps(all_cards, indent=2) + "\n```"
            intro = (
                f"📄 **Resume Analyzed!**\n\n"
                f"🔑 **Skills Detected:** {skills_display}\n"
                f"💼 **Target Roles:** {roles_display}\n"
                f"🔍 **Searched for:** `{search_query}`\n"
                f"⏰ **Filter:** Jobs posted in the **last 7 days** only\n\n"
                f"Found **{len(all_cards)} fresh, open opportunities** matching your profile:\n\n"
            )
            ai_response = intro + cards_block
        else:
            ai_response = (
                f"📄 **Resume Analyzed!**\n\n"
                f"🔑 **Skills Detected:** {skills_display}\n"
                f"💼 **Target Roles:** {roles_display}\n\n"
                f"⚠️ No jobs found in the **last 7 days** for `{search_query}`. "
                f"Ask me to search with a wider range, e.g. *'find python developer jobs from last month'*."
            )

        append_chat_messages(session_id, [
            HumanMessage(content=f"[Resume analyzed — skills: {', '.join(found_skills[:5])}]"),
            AIMessage(content=ai_response)
        ])

        return jsonify({
            "response": ai_response,
            "status": "success",
            "skills_found": found_skills,
            "roles_found": found_roles,
            "search_query": search_query,
            "jobs_count": len(all_cards)
        })

    except Exception as e:
        print(f"Error in /upload-resume: {e}")
        return jsonify({"error": f"Failed to process resume: {str(e)}", "status": "error"}), 500



@app.route("/api/avatar/session", methods=["POST"])
def create_avatar_session():
    """Create a new Akool Live Avatar streaming session."""
    akool_api_key = os.getenv("AKOOL_API_KEY")
    if not akool_api_key or akool_api_key == "your_akool_api_key_here":
        return jsonify({"error": "AKOOL_API_KEY is not set in the .env file"}), 500

    url = "https://openapi.akool.com/api/v3/streamingAvatar/session/create"
    headers = {
        "Authorization": f"Bearer {akool_api_key}",
        "Content-Type": "application/json"
    }
    
    # User should set AKOOL_AVATAR_ID in .env, fallback to a commonly used default format or name
    avatar_id = os.getenv("AKOOL_AVATAR_ID", "default_avatar") 
    
    payload = {
        "avatar_id": avatar_id
    }

    try:
        response = requests.post(url, headers=headers, json=payload, timeout=EXTERNAL_HTTP_TIMEOUT)
        response.raise_for_status()
        data = response.json()
        return jsonify(data)
    except requests.exceptions.HTTPError as e:
        resp_text = e.response.text if e.response is not None else str(e)
        resp_status = e.response.status_code if e.response is not None else 500
        print(f"HTTP Error creating Akool session: {resp_text}")
        return jsonify({"error": "Failed to create avatar session", "details": resp_text}), resp_status
    except Exception as e:
        print(f"Error creating Akool session: {e}")
        return jsonify({"error": str(e)}), 500



# ──────────────────────────────────────────────
# Mock Interview Question Generator
# ──────────────────────────────────────────────
@app.route('/api/interview/generate-questions', methods=['POST'])
def generate_interview_questions():
    try:
        data = request.get_json() or {}
        difficulty = data.get('difficulty', 'medium')
        try:
            count = int(data.get('count', 20))
        except (TypeError, ValueError):
            count = 20
        from mock_interview import generate_real_life_questions, infer_role

        incoming_profile = data.get("profile")
        if isinstance(incoming_profile, dict) and incoming_profile:
            profile = dict(incoming_profile)
        else:
            skills = data.get("skills", []) or []
            role = data.get("role") or (infer_role(skills) if skills else "Software Engineer")
            profile = {
                "role": role,
                "skills": skills,
                "projects": [],
                "experienceHighlights": [],
                "achievements": [],
                "yearsExperience": 0
            }

        questions = generate_real_life_questions(profile, difficulty=difficulty, count=count)
        return jsonify({
            'questions': questions,
            'difficulty': difficulty,
            'count': len(questions),
            'role': profile.get("role", "Software Engineer")
        })
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500



# ──────────────────────────────────────────────
# ATS System - Match Analysis
# ──────────────────────────────────────────────
@app.route('/api/bot/check-ats', methods=['POST'])
def check_ats_score():
    """Detailed ATS analysis between stored resume and a specific job card."""
    data = request.get_json() or {}
    sid = data.get('session_id')
    job_title = data.get('job_title', 'Unknown Role')
    job_desc = data.get('job_description', '')
    
    if not sid:
        return jsonify({"error": "No session ID provided"}), 400
    
    # Try to find user by session_id
    user = db.get_user_by_session_id(sid)
    
    resume_text = ""
    if user:
        resume_text = user.get('resume_text', '')
    
    # If no resume text in DB, try to extract from the PDF file if it exists
    if not resume_text:
        filepath = os.path.join(UPLOAD_FOLDER, f"{sid}.pdf")
        if os.path.exists(filepath):
            try:
                from mock_interview import extract_text_from_pdf
                with open(filepath, 'rb') as f:
                    resume_text = extract_text_from_pdf(io.BytesIO(f.read()))
            except Exception as e:
                print(f"[ATS] Extraction error: {e}")

    if not resume_text:
        return jsonify({"error": "Please upload your resume first to run a detailed ATS check."}), 400

    llm = get_llm()
    if not llm:
        return jsonify({"error": "AI service is currently unavailable."}), 503
    
    prompt = f"""
    Perform a professional, encouraging ATS (Applicant Tracking System) analysis for this internship match.
    Be very specific and actionable.
    
    INTERNSHIP: {job_title}
    DESCRIPTION: {job_desc}
    
    CANDIDATE PROFILE:
    {resume_text[:4000]}
    
    Return ONLY a JSON object with:
    - match_score: (Integer 0-100)
    - match_level: (String: "High Match", "Potential Match", "Gap Detected")
    - missing_keywords: (Array of technical terms missing)
    - strengths: (Array of 3 fit points)
    - improvements: (Array of 3 specific resume editing tips)
    - verdict: (A 2-sentence summary)
    """
    
    try:
        from langchain_core.messages import HumanMessage
        resp = llm.invoke([HumanMessage(content=prompt)])
        # Clean up possible markdown in response
        content = resp.content if isinstance(resp.content, str) else str(resp.content)
        raw_content = content.strip()
        if "```json" in raw_content:
            raw_content = raw_content.split("```json")[1].split("```")[0].strip()
        elif "```" in raw_content:
             raw_content = raw_content.split("```")[1].split("```")[0].strip()
        
        analysis = json.loads(raw_content)
        return jsonify(analysis)
    except Exception as e:
        print(f"[ATS] Error: {e}")
        import traceback; traceback.print_exc()
        return jsonify({"error": "Failed to generate ATS analysis. Please try again."}), 500

@app.route('/api/fast-extract-skills', methods=['POST'])
def fast_extract_skills():
    if 'resume' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    try:
        file = request.files['resume']
        from mock_interview import extract_text_from_pdf, extract_resume_profile
        text = extract_text_from_pdf(file.stream)
        if not text or len(text) < 30:
            return jsonify({"error": "Could not extract readable text from the resume PDF"}), 400
        profile = extract_resume_profile(text)
        return jsonify({
            "extractedSkills": profile.get("skills", []),
            "role": profile.get("role", "Software Engineer"),
            "profile": profile
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ──────────────────────────────────────────────
# Run the server
# ──────────────────────────────────────────────
if __name__ == "__main__":
    host = os.getenv("HOST", "0.0.0.0")
    port = int(os.getenv("PORT", "5000"))
    debug_flag = os.getenv("DEBUG", "").strip().lower() in {"1", "true", "yes"}
    env_name = os.getenv("FLASK_ENV", "").strip().lower()
    is_dev = debug_flag or env_name == "development"

    print("[START] AI Internship Finder Agent is running!")
    print(f"[INFO] Listening on http://{host}:{port}")
    print("")

    if not is_dev:
        try:
            from waitress import serve
            threads = max(4, int(os.getenv("WAITRESS_THREADS", "16")))
            connection_limit = max(threads * 4, int(os.getenv("WAITRESS_CONNECTION_LIMIT", str(threads * 8))))
            print(f"[PROD] Waitress enabled with threads={threads}, connection_limit={connection_limit}")
            serve(app, host=host, port=port, threads=threads, connection_limit=connection_limit)
        except Exception as ex:
            print(f"[WARN] Waitress unavailable ({ex}). Falling back to Flask threaded server.")
            app.run(debug=False, host=host, port=port, threaded=True, use_reloader=False)
    else:
        print("[DEV] Running Flask debug server")
        app.run(debug=True, host=host, port=port, threaded=True, use_reloader=True)
